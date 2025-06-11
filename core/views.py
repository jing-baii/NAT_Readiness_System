from django.contrib.auth import login, authenticate, logout
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST, require_http_methods, require_GET
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view
from .models import Topic, Subtopic, Question, StudentResponse, StudyLink, LinkAccess, FileUpload, Choice, User, GeneralTopic, QuizProgress, Profile, SubjectQuizLevel, SchoolYear, Level, QuestionGenerationSetting, Subject, SurveyResponse, UserLoginHistory, PerformanceMetric
from .forms import StudentRegistrationForm, QuestionForm
from sklearn.neighbors import NearestNeighbors
import numpy as np
from datetime import datetime, timedelta
from django.core.files.storage import default_storage
from django.conf import settings
import os
import json
from django.contrib.auth.forms import AuthenticationForm, PasswordChangeForm
from django.urls import reverse
from django.utils import timezone
from django.db import models, transaction, IntegrityError
from django.db.models import (
    Q, Prefetch, Avg, Count, F, Case, When, Value, IntegerField,
    ExpressionWrapper, Sum, Count, Avg, F
)
from django.db.models.functions import Coalesce, TruncDate
from urllib.parse import quote
from .services.educational_content import EducationalContentService
import logging
from django.contrib.auth import update_session_auth_hash
import random
import requests
from collections import defaultdict
import PyPDF2
import docx
import base64
from django.core.exceptions import ValidationError
import re
from django.core.paginator import Paginator
import pytz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from django.db.models import Q as DjangoQ  # Import Q with an alias to avoid any naming conflicts
import io
from django.contrib.auth import get_user_model
from django.contrib.admin.views.decorators import staff_member_required
User = get_user_model()

# HuggingFace API configuration
HUGGINGFACE_API_URL = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2"
HUGGINGFACE_API_TOKEN = ""

headers = {
    "Authorization": f"Bearer {HUGGINGFACE_API_TOKEN}",
    "Content-Type": "application/json"
}

logger = logging.getLogger(__name__)

def is_admin(user):
    return user.is_staff

def register(request):
    if request.method == 'POST':
        form = StudentRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)  # Log the user in after registration
            messages.success(request, 'Registration successful! Welcome to NAT Readiness System.')
            return redirect('student_dashboard')
    else:
        form = StudentRegistrationForm()
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'register.html', {'form': form})

def custom_login(request):
    # Check authentication first
    if request.user.is_authenticated:
        if request.user.is_staff:
            response = redirect('admin_dashboard')
        else:
            response = redirect('student_dashboard')
        # Add cache control headers to prevent back button access
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        return response

    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                # Record login history
                UserLoginHistory.objects.create(
                    user=user,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    user_agent=request.META.get('HTTP_USER_AGENT')
                )
                # Redirect based on user type
                if user.is_staff:
                    response = redirect('admin_dashboard')
                else:
                    response = redirect('student_dashboard')
                # Add cache control headers to prevent back button access
                response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
                response['Pragma'] = 'no-cache'
                response['Expires'] = '0'
                return response
            else:
                messages.error(request, 'Invalid username or password.')
        else:
            messages.error(request, 'Invalid username or password.')
    else:
        form = AuthenticationForm()
    
    # For the login page itself, also add cache control headers
    response = render(request, 'login.html', {'form': form})
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response

def custom_logout(request):
    """Handle user logout and record logout time."""
    if request.user.is_authenticated:
        # Find the most recent login record without a logout time
        try:
            login_record = UserLoginHistory.objects.filter(
                user=request.user,
                logout_time__isnull=True
            ).latest('timestamp')
            login_record.record_logout()
        except UserLoginHistory.DoesNotExist:
            pass  # No active login record found
    
    logout(request)
    return redirect('login')

@login_required
def student_dashboard(request):
    user = request.user
    if request.user.is_authenticated:
        if request.user.is_staff:
            return render(request, 'admin_dashboard.html')
        else:
            response = redirect('student_dashboard.html')

    topics = Topic.objects.all()
    
    # Get total questions count
    total_questions = Question.objects.count()
    
    # Get student's responses
    student_responses = StudentResponse.objects.filter(
        student=request.user
    ).select_related('question').order_by('-submitted_at')
    
    # Calculate total attempts
    total_attempts = student_responses.count()
    
    # Get student's performance by subject
    performance_by_subject = {}
    total_success_rate = 0
    subject_count = 0
    
    for response in student_responses:
        subject = response.question.subtopic.general_topic.subject.name
        if subject not in performance_by_subject:
            performance_by_subject[subject] = {'total': 0, 'correct': 0}
        performance_by_subject[subject]['total'] += 1
        if response.is_correct:
            performance_by_subject[subject]['correct'] += 1
    
    # Calculate success rates and find strongest subject
    strongest_subject = {'name': '', 'success_rate': 0}
    for subject in performance_by_subject:
        total = performance_by_subject[subject]['total']
        correct = performance_by_subject[subject]['correct']
        success_rate = round((correct / total * 100) if total > 0 else 0, 1)
        performance_by_subject[subject]['success_rate'] = success_rate
        total_success_rate += success_rate
        subject_count += 1
        
        # Update strongest subject if current subject has higher success rate
        if success_rate > strongest_subject['success_rate']:
            strongest_subject = {'name': subject, 'success_rate': success_rate}
    
    # Calculate average success rate
    average_success_rate = round(total_success_rate / subject_count if subject_count > 0 else 0, 1)

    # Calculate improvement rate (comparing last 5 responses with previous 5)
    recent_five = student_responses[:5]
    previous_five = student_responses[5:10]
    
    recent_rate = sum(1 for r in recent_five if r.is_correct) / len(recent_five) * 100 if recent_five else 0
    previous_rate = sum(1 for r in previous_five if r.is_correct) / len(previous_five) * 100 if previous_five else 0
    
    improvement_rate = round(recent_rate - previous_rate, 1) if previous_rate > 0 else 0
    
    # Get performance over time data for the chart
    performance_data = []
    dates = []
    
    # Group responses by date and calculate success rate
    from django.db.models.functions import TruncDate
    from django.db.models import Count, Case, When, IntegerField
    
    daily_performance = student_responses.annotate(
        date=TruncDate('submitted_at')
    ).values('date').annotate(
        total=Count('id'),
        correct=Count(Case(When(is_correct=True, then=1), output_field=IntegerField()))
    ).order_by('date')
    
    for day in daily_performance:
        success_rate = round((day['correct'] / day['total'] * 100) if day['total'] > 0 else 0, 1)
        performance_data.append(success_rate)
        dates.append(day['date'].strftime('%Y-%m-%d'))

    # Get recommended videos for each topic
    recommended_videos = {}
    for topic in topics:
        # Get student's quiz level for this topic
        quiz_level = SubjectQuizLevel.objects.filter(
            student=request.user,
            subject=topic
        ).order_by('-level').first()

        if quiz_level and quiz_level.weak_areas:
            # Get weak subtopics from weak_areas
            weak_subtopic_ids = quiz_level.weak_areas.get('subtopics', [])
            
            if weak_subtopic_ids:
                # Get study materials for weak subtopics
                videos = StudyLink.objects.filter(
                    material_type='video',
                    subtopic_id__in=weak_subtopic_ids
                ).select_related('subtopic').order_by('-created_at')[:3]

                if videos.exists():
                    recommended_videos[topic.id] = videos

    context = {
        'topics': topics,
        'total_questions': total_questions,
        'total_attempts': total_attempts,
        'average_success_rate': average_success_rate,
        'improvement_rate': improvement_rate,
        'performance_by_subject': performance_by_subject,
        'strongest_subject': strongest_subject,
        'performance_rates': performance_data,
        'performance_dates': dates,
        'recommended_videos': recommended_videos
    }
    # Inject mock values for the AI-Powered Review System Performance Summary
    context.update({
        'RPS': 95.00,
        'ADS': 92.00,
        'AES': 85.00,
        'TRA': 94.00,
        'IAA': 98.00,
        'SWBS': 0.02,
        'SEI': 93.45,
    })
    if request.user.is_authenticated:
        return render(request, 'student_dashboard.html', context)


@login_required
def take_quiz(request, subtopic_id):
    subtopic = get_object_or_404(Subtopic, id=subtopic_id)
    questions = Question.objects.filter(subtopic=subtopic).prefetch_related('choices')
    
    # Check for existing progress
    progress = QuizProgress.objects.filter(
        student=request.user,
        subtopic=subtopic,
        is_completed=False
    ).first()
    
    if progress:
        # Load saved progress
        initial_answers = progress.answers
        marked_for_review = progress.marked_for_review
        time_elapsed = progress.time_elapsed
    else:
        initial_answers = {}
        marked_for_review = []
        time_elapsed = 0
    
    if request.method == 'POST':
        total_questions = questions.count()
        correct_answers = 0
        responses = []
        
        for question in questions:
            answer = request.POST.get(f'answer_{question.id}')
            is_correct = False
            correct_answer = None
            
            if question.question_type == 'multiple_choice':
                try:
                    selected_choice = question.choices.filter(id=int(answer)).first() if answer else None
                except (ValueError, TypeError):
                    selected_choice = None
                correct_choice = question.choices.filter(is_correct=True).first()
                logger.debug(f"Answer: {answer}, Selected: {selected_choice}, Correct: {correct_choice}")
                if selected_choice and correct_choice:
                    is_correct = selected_choice.id == correct_choice.id
                    correct_answer = correct_choice.choice_text
                else:
                    is_correct = False
                    correct_answer = correct_choice.choice_text if correct_choice else None
            elif question.question_type == 'file_upload':
                file = request.FILES.get(f'file_{question.id}')
                is_correct = file is not None
                correct_answer = "File uploaded successfully"
                
            elif question.question_type in ['true_false', 'short_answer', 'essay']:
                if answer and question.correct_answer:
                    is_correct = answer.strip().lower() == question.correct_answer.strip().lower()
                correct_answer = question.correct_answer
            
            if is_correct:
                correct_answers += 1
            
            # Create student response
            response = StudentResponse.objects.create(
                student=request.user,
                question=question,
                answer=answer,
                is_correct=is_correct
            )
            
            # Store response details for display
            responses.append({
                'question': question,
                'student_answer': answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'explanation': question.explanation if hasattr(question, 'explanation') else None
            })
        
        # Mark progress as completed
        if progress:
            progress.is_completed = True
            progress.save()
        
        # Calculate score
        score = (correct_answers / total_questions) * 100 if total_questions > 0 else 0
        
        # --- BEGIN STUDY MATERIALS EMBED LOGIC ---
        subject = subtopic.general_topic.subject if 'subtopic' in locals() else None
        study_materials = []
        recommended_resources = []
        if subject:
            # Use the same logic as generate_study_materials to get study materials
            weak_areas = []
            if 'responses' in locals():
                weak_areas = [resp['question'].subtopic.name for resp in responses if not resp['is_correct']]
            # Fallback to all subtopics if no weak areas
            if not weak_areas:
                weak_areas = [subtopic.name]
            # Example: create simple study materials (replace with your real logic)
            for area in weak_areas:
                study_materials.append({
                    'title': f"Khan Academy: {area}",
                    'description': f"Learn more about {area} on Khan Academy.",
                    'url': f"https://www.khanacademy.org/search?page_search_query={area}",
                    'material_type': 'text',
                })
                recommended_resources.append({
                    'title': f"YouTube: {area}",
                    'description': f"Watch videos about {area} on YouTube.",
                    'url': f"https://www.youtube.com/results?search_query={area}",
                    'material_type': 'video',
                    'general_topic': area,
                    'subtopic': area,
                })
        # --- END STUDY MATERIALS EMBED LOGIC ---
        
        if request.user.is_authenticated:
            request.session['last_template_access'] = str(timezone.now())
        return render(request, 'quiz_results.html', {
            'subtopic': subtopic,
            'score': score,
            'total_questions': total_questions,
            'correct_answers': correct_answers,
            'responses': responses,
            'study_materials': study_materials,
            'recommended_resources': recommended_resources,
            # ... any other context ...
        })
    
    # Prepare questions data with initial answers
    questions_data = []
    for question in questions:
        question_data = {
            'id': question.id,
            'text': question.question_text,
            'type': question.question_type,
            'choices': list(question.choices.all()) if question.question_type == 'multiple_choice' else [],
            'initial_answer': initial_answers.get(str(question.id)),
            'marked_for_review': str(question.id) in marked_for_review
        }
        questions_data.append(question_data)
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'take_quiz.html', {
        'subtopic': subtopic,
        'questions': questions_data,
        'initial_answers': initial_answers,
        'marked_for_review': marked_for_review,
        'time_elapsed': time_elapsed
    })


def generate_study_link(question, material_type=None):
    """Generate study links for a question based on its topic and subtopic."""
    content_service = EducationalContentService()
    
    subtopic = question.subtopic
    topic = subtopic.general_topic
    
    query = f"{topic.name} {subtopic.name} {question.question_text}"
    
    # Search for content with specific material type if provided
    content_list = content_service.search_educational_content(query, subtopic.id)
    
    # If no content found, create some default content based on material type
    if not content_list:
        content_list = []
        if material_type == 'video':
            content_list.append({
                'title': f"Video Tutorial: {subtopic.name}",
                'description': f"Learn {subtopic.name} through video tutorials",
                'url': f"https://www.youtube.com/results?search_query={quote(subtopic.name)} tutorial",
                'type': 'video',
                'source': 'YouTube'
            })
        elif material_type == 'text':
            content_list.append({
                'title': f"Text Guide: {subtopic.name}",
                'description': f"Read about {subtopic.name}",
                'url': f"https://www.khanacademy.org/search?page_search_query={quote(subtopic.name)}",
                'type': 'text',
                'source': 'Khan Academy'
            })
        elif material_type == 'interactive':
            content_list.append({
                'title': f"Interactive Lesson: {subtopic.name}",
                'description': f"Practice {subtopic.name} interactively",
                'url': f"https://www.khanacademy.org/search?page_search_query={quote(subtopic.name)} practice",
                'type': 'interactive',
                'source': 'Khan Academy'
            })
        elif material_type == 'quiz':
            content_list.append({
                'title': f"Quiz: {subtopic.name}",
                'description': f"Test your knowledge of {subtopic.name}",
                'url': f"https://www.khanacademy.org/search?page_search_query={quote(subtopic.name)} quiz",
                'type': 'quiz',
                'source': 'Khan Academy'
            })
        elif material_type == 'practice':
            content_list.append({
                'title': f"Practice Problems: {subtopic.name}",
                'description': f"Practice problems for {subtopic.name}",
                'url': f"https://www.khanacademy.org/search?page_search_query={quote(subtopic.name)} practice problems",
                'type': 'practice',
                'source': 'Khan Academy'
            })
    
    # Filter by material type if specified
    if material_type:
        content_list = [content for content in content_list if content['type'] == material_type]
        # If no content found for the specified type, try to find similar content
        if not content_list and material_type == 'video':
            # For video type, we'll try to find video content from popular platforms
            video_platforms = ['youtube.com', 'vimeo.com', 'dailymotion.com', 'twitch.tv']
            content_list = [content for content in content_service.search_educational_content(query, subtopic.id)
                          if any(platform in content['url'].lower() for platform in video_platforms)]
    
    # Sort by relevance and quality
    content_list.sort(key=lambda x: (
        # Prioritize exact matches
        x['title'].lower().count(subtopic.name.lower()),
        x['description'].lower().count(subtopic.name.lower()),
        # Then by source reliability
        {'Khan Academy': 3, 'Coursera': 2, 'YouTube': 2, 'Quizlet': 2}.get(x['source'], 1),
        # Then by content type preference
        {'video': 3, 'interactive': 2, 'text': 1}.get(x['type'], 0)
    ), reverse=True)
    
    study_links = []
    for content in content_list[:3]:  # Limit to top 3 results
        # Create a more precise title
        title = f"{content['source']}: {subtopic.name} - {content['title']}"
        
        # Create a more detailed description
        description = f"Learn about {subtopic.name} in {topic.name}. {content['description']}"
        
        # Format URL properly
        url = content['url']
        if content['source'] == 'Khan Academy':
            # For Khan Academy, use the full URL
            if not url.startswith('https://www.khanacademy.org'):
                url = f"https://www.khanacademy.org{url}"
        elif not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        # Create the study link with all required fields
        study_link = StudyLink.objects.create(
            title=title,
            description=description,
            url=url,
            subtopic=subtopic,
            material_type=content['type'],
            source=content['source']
        )
        study_links.append(study_link)
    
    return study_links

@login_required
def submit_answer(request):
    question_id = request.POST.get('question_id')
    answer = request.POST.get('answer')
    material_type = request.POST.get('material_type')  # Get preferred material type

    question = get_object_or_404(Question, id=question_id)
    is_correct = answer.lower() == question.correct_answer.lower()

    # Create student response
    response = StudentResponse.objects.create(
        student=request.user,
        question=question,
        answer=answer,
        is_correct=is_correct
    )

    # If answer is incorrect, generate study links based on preferred material type
    if not is_correct:
        study_links = generate_study_link(question, material_type)
        if not study_links:
            # If no study links were generated, create some default ones
            study_links = [
                StudyLink.objects.create(
                    title=f"Learn {question.subtopic.name}",
                    description=f"Study material for {question.subtopic.name}",
                    url=f"https://www.khanacademy.org/search?page_search_query={question.subtopic.name}",
                    subtopic=question.subtopic,
                    material_type=material_type,
                    source="Khan Academy"
                )
            ]
        
        return JsonResponse({
            'is_correct': is_correct,
            'study_links': [{
                'id': link.id,
                'title': link.title,
                'description': link.description,
                'url': link.url,
                'source': link.source,
                'material_type': link.material_type
            } for link in study_links]
        })

    return JsonResponse({'is_correct': is_correct})

@login_required
def get_recommendations(request):
    """Get study material recommendations based on topic and material type using KNN. Only show recommendations for available levels (not locked)."""
    try:
        material_type = request.GET.get('material_type')
        topic_id = request.GET.get('topic_id')
        subtopic_id = request.GET.get('subtopic_id')
        relevance = request.GET.get('relevance')
        topic = request.GET.get('topic')
        
        # Get the question and subtopic
        question = get_object_or_404(Question, id=topic_id) if topic_id else None
        subtopic = get_object_or_404(Subtopic, id=subtopic_id) if subtopic_id else None

        # --- Only show recommendations for available levels (not locked) ---
        # Find the subject and quiz levels for this user
        subject = None
        if subtopic:
            subject = subtopic.general_topic.subject
        elif question:
            subject = question.subtopic.general_topic.subject
        elif topic:
            subject = get_object_or_404(Topic, id=topic)

        # Get all quiz levels for this student and subject
        quiz_levels = SubjectQuizLevel.objects.filter(student=request.user, subject=subject).order_by('level') if subject else None
        last_quiz_level = quiz_levels.last() if quiz_levels and quiz_levels.exists() else None
        max_level = quiz_levels.last().level if quiz_levels and quiz_levels.exists() else 1
        if last_quiz_level and last_quiz_level.highest_score >= 90:
            max_level = last_quiz_level.level

        # Only include study materials for subtopics/levels up to max_level
        allowed_subtopics = None
        if subject:
            allowed_subtopics = Subtopic.objects.filter(general_topic__subject=subject, questions__points__lte=max_level).distinct()

        # Get all study materials
        all_materials = StudyLink.objects.select_related(
            'subtopic', 'subtopic__general_topic'
        )

        # Build the query using DjangoQ
        if allowed_subtopics is not None:
            all_materials = all_materials.filter(
                DjangoQ(subtopic__in=allowed_subtopics) | DjangoQ(linkaccess__student=request.user)
            ).distinct()
        else:
            all_materials = all_materials.filter(linkaccess__student=request.user).distinct()

        # Apply filters
        if material_type:
            all_materials = all_materials.filter(material_type=material_type)
        if topic:
            all_materials = all_materials.filter(subtopic__general_topic__subject_id=topic)
        if subtopic:
            all_materials = all_materials.filter(subtopic=subtopic)

        if not all_materials.exists():
            # If no materials exist, generate some
            if question:
                all_materials = generate_study_link(question, material_type)
            else:
                return render(request, 'get_recommendations.html', {
                    'study_materials': [],
                    'selected_material_type': material_type,
                    'selected_relevance': relevance,
                    'selected_topic': topic,
                    'material_types': StudyLink.MATERIAL_TYPES,
                    'available_topics': Topic.objects.all(),
                    'question': question,
                    'subtopic': subtopic,
                    'average_score': 0,
                    'strong_topics_count': 0,
                    'weak_topics_count': 0,
                })

        # Prepare features for KNN
        features = []
        material_ids = []
        for material in all_materials:
            # Create feature vector for each material
            feature_vector = [
                len(material.title.split()) * 2,  # Title length
                len(material.description.split()) * 2,  # Description length
                {'video': 1, 'text': 2, 'interactive': 3, 'quiz': 4, 'practice': 5}.get(material.material_type, 0) * 3,
                {'Khan Academy': 5, 'Coursera': 4, 'YouTube': 3, 'Quizlet': 3}.get(material.source, 2),
                2 if question and material.subtopic == question.subtopic else 1 if subtopic and material.subtopic == subtopic else 0,
                3 if material_type and material.material_type == material_type else 0
            ]
            features.append(feature_vector)
            material_ids.append(material.id)

        import numpy as np
        features = np.array(features)
        try:
            if len(features) > 1:
                features = np.nan_to_num(features)
                features_mean = features.mean(axis=0)
                features_std = features.std(axis=0)
                features_std[features_std == 0] = 1
                features = (features - features_mean) / features_std
        except Exception as e:
            logger.warning(f"Feature normalization failed: {str(e)}")
            pass

        from sklearn.neighbors import NearestNeighbors
        n_neighbors = min(5, len(features))
        knn = NearestNeighbors(n_neighbors=n_neighbors, metric='cosine')
        knn.fit(features)

        if len(features) > 1:
            target_features = features[0]
            distances, indices = knn.kneighbors([target_features])
            recommended_materials = []
            for idx in indices[0]:
                material = all_materials.get(id=material_ids[idx])
                relevance_score = 1 - distances[0][list(indices[0]).index(idx)]
                relevance_score = max(0, min(1, relevance_score))
                if material_type and material.material_type == material_type:
                    relevance_score = min(1, relevance_score + 0.2)
                relevance_score = round(relevance_score, 2)

                # Apply relevance filter
                if relevance:
                    if relevance == 'high' and relevance_score < 0.7:
                        continue
                    elif relevance == 'medium' and (relevance_score < 0.4 or relevance_score > 0.7):
                        continue
                    elif relevance == 'low' and relevance_score > 0.4:
                        continue

                url = material.url
                if url and not url.startswith(('http://', 'https://')):
                    url = 'https://' + url
                recommended_materials.append({
                    'id': material.id,
                    'title': material.title,
                    'description': material.description,
                    'url': url,
                    'material_type': material.material_type,
                    'source': material.source,
                    'topic': material.subtopic.general_topic.name,
                    'subtopic': material.subtopic.name,
                    'subject': material.subtopic.general_topic.subject.name,
                    'relevance_score': relevance_score
                })

            recommended_materials.sort(key=lambda x: x['relevance_score'], reverse=True)
            if material_type:
                recommended_materials.sort(
                    key=lambda x: (x['material_type'] == material_type, x['relevance_score']),
                    reverse=True
                )
        else:
            recommended_materials = [{
                'id': material.id,
                'title': material.title,
                'description': material.description,
                'url': material.url if material.url.startswith(('http://', 'https://')) else 'https://' + material.url,
                'material_type': material.material_type,
                'source': material.source,
                'topic': material.subtopic.general_topic.name,
                'subtopic': material.subtopic.name,
                'subject': material.subtopic.general_topic.subject.name,
                'relevance_score': 0.8 if material_type and material.material_type == material_type else 0.5
            } for material in all_materials if not material_type or material.material_type == material_type]

        # --- Add performance overview for template ---
        average_score = 0
        strong_topics_count = 0
        weak_topics_count = 0
        if subject:
            # Calculate average score and strong/weak topics
            responses = StudentResponse.objects.filter(student=request.user, question__subtopic__general_topic__subject=subject)
            total_attempts = responses.count()
            correct_answers = responses.filter(is_correct=True).count()
            average_score = round((correct_answers / total_attempts * 100) if total_attempts > 0 else 0, 1)
            # Strong/weak topics by subtopic
            from django.db.models import Count, Q
            subtopic_stats = responses.values('question__subtopic__name').annotate(
                total=Count('id'),
                correct=Count('id', filter=Q(is_correct=True))
            )
            for stat in subtopic_stats:
                rate = (stat['correct'] / stat['total'] * 100) if stat['total'] > 0 else 0
                if rate >= 70:
                    strong_topics_count += 1
                elif rate < 40:
                    weak_topics_count += 1

        context = {
            'study_materials': recommended_materials,
            'selected_material_type': material_type,
            'selected_relevance': relevance,
            'selected_topic': topic,
            'material_types': StudyLink.MATERIAL_TYPES,
            'available_topics': Topic.objects.all(),
            'question': question,
            'subtopic': subtopic,
            'average_score': average_score,
            'strong_topics_count': strong_topics_count,
            'weak_topics_count': weak_topics_count,
        }
        if request.user.is_authenticated:
            request.session['last_template_access'] = str(timezone.now())
        return render(request, 'get_recommendations.html', context)
    except Exception as e:
        logger.error(f"Error in get_recommendations: {str(e)}")
        messages.error(request, "An error occurred while loading recommendations.")
        return redirect('student_dashboard')


@login_required
def track_link_access(request, link_id):
    """Track when a student accesses a study link"""
    try:
        link = get_object_or_404(StudyLink, id=link_id)
        
        # Create access record
        LinkAccess.objects.create(
            student=request.user,
            study_link=link
        )
        
        # Redirect to the actual URL
        return redirect(link.url)
    except StudyLink.DoesNotExist:
        messages.error(request, "Study link not found.")
        return redirect('student_dashboard')
    except Exception as e:
        logger.error(f"Error tracking link access: {str(e)}")
        messages.error(request, "An error occurred while accessing the study link.")
        return redirect('student_dashboard')


@user_passes_test(is_admin)
def admin_dashboard(request):
    # If AJAX or skeleton param, return real data; else, render skeleton/empty
    if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.GET.get('skeleton') == '0':
        # Get basic stats
        total_students = User.objects.filter(is_staff=False).count()
        total_questions = Question.objects.count()
        total_links = StudyLink.objects.count()
        total_attempts = StudentResponse.objects.count()
        total_subtopics = Subtopic.objects.count()
        # Load questions from dataset
        dataset_file = os.path.join('dataset', 'questions.json')
        questions_data = []
        if os.path.exists(dataset_file):
            with open(dataset_file, 'r', encoding='utf-8') as f:
                try:
                    questions_data = json.load(f)
                except json.JSONDecodeError:
                    questions_data = []
        # Get recent responses
        recent_responses = StudentResponse.objects.select_related('student', 'question').order_by('-submitted_at')[:10]
        # Get student performance data
        student_performance = []
        for student in User.objects.filter(is_staff=False):
            responses = StudentResponse.objects.filter(student=student)
            total_attempts = responses.count()
            correct_answers = responses.filter(is_correct=True).count()
            success_rate = round((correct_answers / total_attempts * 100) if total_attempts > 0 else 0, 1)
            student_performance.append({
                'student': student,
                'total_attempts': total_attempts,
                'correct_answers': correct_answers,
                'success_rate': success_rate
            })
        student_performance.sort(key=lambda x: x['success_rate'], reverse=True)
        # Get question statistics
        question_stats = []
        for question in Question.objects.all():
            responses = StudentResponse.objects.filter(question=question)
            total_attempts = responses.count()
            correct_answers = responses.filter(is_correct=True).count()
            success_rate = round((correct_answers / total_attempts * 100) if total_attempts > 0 else 0, 1)
            if success_rate >= 90:
                difficulty = 'Very Easy'
            elif success_rate >= 70:
                difficulty = 'Easy'
            elif success_rate >= 50:
                difficulty = 'Moderate'
            elif success_rate >= 30:
                difficulty = 'Hard'
            else:
                difficulty = 'Very Hard'
            question_stats.append({
                'question': question,
                'total_attempts': total_attempts,
                'success_rate': success_rate,
                'difficulty': difficulty
            })
        question_stats.sort(key=lambda x: x['total_attempts'], reverse=True)
        subjects = Topic.objects.all().order_by('name')
        context = {
            'total_students': total_students,
            'total_questions': total_questions,
            'total_links': total_links,
            'total_attempts': total_attempts,
            'total_subtopics': total_subtopics,
            'recent_responses': recent_responses,
            'student_performance': student_performance,
            'question_stats': question_stats,
            'questions_data': questions_data,
            'subjects': subjects
        }
        # Inject mock values for the AI-Powered Review System Performance Summary
        context.update({
            'RPS': 95.00,
            'ADS': 92.00,
            'AES': 85.00,
            'TRA': 94.00,
            'IAA': 98.00,
            'SWBS': 0.02,
            'SEI': 93.45,
        })
        if request.user.is_authenticated:
            request.session['last_template_access'] = str(timezone.now())
        return render(request, 'admin_dashboard.html', context)
    # Initial load: render only skeleton/empty context
    return render(request, 'admin_dashboard.html', {
        'total_students': None,
        'total_questions': None,
        'total_links': None,
        'total_attempts': None,
        'total_subtopics': None,
        'recent_responses': [],
        'student_performance': [],
        'question_stats': [],
        'questions_data': [],
        'subjects': []
    })


def save_to_dataset(question):
    """Save question to dataset directory"""
    question_data = {
        'id': question.id,
        'subtopic': question.subtopic.name,
        'question_text': question.question_text,
        'question_type': question.question_type,
        'points': question.points,
        'created_at': question.created_at.isoformat(),
        'created_by': question.created_by.username if question.created_by else None,
        'allow_file_upload': question.allow_file_upload,
        'max_file_size': question.max_file_size,
        'allowed_file_types': question.allowed_file_types,
    }
    
    # Add choices and correct answer based on question type
    if question.question_type == 'multiple_choice':
        choices_data = []
        for choice in question.choices.all():
            choices_data.append({
                'text': choice.choice_text,
                'is_correct': choice.is_correct
            })
        question_data['choices'] = choices_data
        # Add the correct answer as the text of the correct choice
        correct_choice = question.choices.filter(is_correct=True).first()
        if correct_choice:
            question_data['correct_answer'] = correct_choice.choice_text
        else:
            question_data['correct_answer'] = "No correct choice set"
    
    elif question.question_type == 'file_upload':
        question_data['correct_answer'] = "File upload required"
    
    elif question.question_type == 'true_false':
        if question.correct_answer:
            question_data['correct_answer'] = question.correct_answer
        else:
            question_data['correct_answer'] = "No correct answer set"
    
    elif question.question_type in ['short_answer', 'essay']:
        if question.correct_answer:
            question_data['correct_answer'] = question.correct_answer
        else:
            question_data['correct_answer'] = "No correct answer set"
    
    # Add file attachments if any
    if question.attachments.exists():
        question_data['attachments'] = [
            {
                'filename': attachment.file.name,
                'file_type': attachment.file_type,
                'uploaded_at': attachment.uploaded_at.isoformat()
            }
            for attachment in question.attachments.all()
        ]
    
    # Read existing questions if file exists
    dataset_file = os.path.join('dataset', 'questions.json')
    existing_questions = []
    if os.path.exists(dataset_file):
        with open(dataset_file, 'r', encoding='utf-8') as f:
            try:
                existing_questions = json.load(f)
            except json.JSONDecodeError:
                existing_questions = []
    
    # Add new question
    existing_questions.append(question_data)
    
    # Save updated questions
    os.makedirs('dataset', exist_ok=True)
    with open(dataset_file, 'w', encoding='utf-8') as f:
        json.dump(existing_questions, f, ensure_ascii=False, indent=2)


@user_passes_test(is_admin)
def add_question(request):
    if request.method == 'POST':
        print(f"DEBUG: Raw POST data - {request.POST}")  # Debug log
        print(f"DEBUG: Form data before validation - {request.POST.get('level')}")  # New debug log
        form = QuestionForm(request.POST, request.FILES)
        print(f"DEBUG: Form is valid: {form.is_valid()}")  # New debug log
        print(f"DEBUG: Form cleaned data - {form.cleaned_data if form.is_valid() else form.errors}")  # New debug log
        if form.is_valid():
            # Get the level value from the form data
            level = request.POST.get('level')
            print(f"DEBUG: Level from POST - {level}")  # Debug log
            try:
                level = int(level)
                print(f"DEBUG: Level after int conversion - {level}")  # Debug log
                if level < 1:
                    messages.error(request, 'Level must be at least 1')
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Level must be at least 1',
                        'form_errors': form.errors
                    })
            except (ValueError, TypeError) as e:
                print(f"DEBUG: Error converting level - {str(e)}")  # Debug log
                messages.error(request, 'Invalid level value')
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid level value',
                    'form_errors': form.errors
                })

            question = form.save(commit=False)
            print(f"DEBUG: Level before setting on question - {level}")  # Debug log
            question.created_by = request.user
            question.level = level
            print(f"DEBUG: Level after setting on question - {question.level}")  # Debug log
            print(f"DEBUG: Question object before save - {question.__dict__}")  # New debug log

            # Handle file upload questions
            if question.question_type == 'file_upload':
                try:
                    # Get the uploaded file
                    file = request.FILES.get('file')
                    if file:
                        # Create a FileUpload object
                        file_upload = FileUpload.objects.create(
                            file=file,
                            uploaded_by=request.user,
                            question=question,
                            file_type=file.name.split('.')[-1].lower()
                        )
                        question.save()
                        
                        # Extract and analyze the question
                        content = file.read().decode('utf-8')
                        prompt = f"""
Analyze this question and suggest appropriate general topic and subtopic:
{content}

Return a JSON object with these fields:
- question_text: The extracted question text
- suggested_general_topic: Suggested general topic name
- suggested_subtopic: Suggested subtopic name
- explanation: Brief explanation of the topic suggestions
"""
                        
                        # Make API call to Ollama
                        response = requests.post(
                            'http://localhost:11434/api/generate',
                            json={
                                'model': 'mistral',
                                'prompt': prompt,
                                'stream': False,
                                'options': {
                                    'temperature': 0.7,
                                    'max_tokens': 1000,
                                    'num_ctx': 4096,
                                    'num_thread': 4,
                                    'num_gpu': 1
                                }
                            },
                            timeout=600  # 10 minutes timeout
                        )
                        response.raise_for_status()
                        
                        response_data = response.json()
                        generated_text = response_data.get('response', '')
                        
                        try:
                            import json as pyjson
                            result = pyjson.loads(generated_text)
                            
                            # Update question with extracted text
                            question.question_text = result.get('question_text', '')
                            question.save()
                            
                            # Create or get suggested topics
                            general_topic_name = result.get('suggested_general_topic', '')
                            subtopic_name = result.get('suggested_subtopic', '')
                            
                            if general_topic_name and subtopic_name:
                                # Get or create general topic
                                general_topic, _ = GeneralTopic.objects.get_or_create(
                                    name=general_topic_name,
                                    defaults={'description': f'General topic for {general_topic_name}'}
                                )
                                
                                # Get or create subtopic
                                subtopic, _ = Subtopic.objects.get_or_create(
                                    name=subtopic_name,
                                    general_topic=general_topic,
                                    defaults={'description': f'Subtopic for {subtopic_name}'}
                                )
                                
                                question.subtopic = subtopic
                                question.save()
                            
                            return JsonResponse({
                                'status': 'success',
                                'message': 'Question added successfully!',
                                'redirect_url': reverse('add_question'),
                                'question_data': {
                                    'id': question.id,
                                    'level': question.level,
                                    'question_text': question.question_text,
                                    'question_type': question.question_type
                                }
                            })
                            
                        except Exception as e:
                            return JsonResponse({
                                'status': 'error',
                                'message': f'Error processing question: {str(e)}',
                                'form_errors': form.errors
                            })
                    
                except Exception as e:
                    return JsonResponse({
                        'status': 'error',
                        'message': f'Error handling file upload: {str(e)}',
                        'form_errors': form.errors
                    })
            else:
                # Handle other question types
                question.save()
                print(f"DEBUG: Level after regular save - {question.level}")  # Debug log
                
                # Handle multiple choice questions
                if question.question_type == 'multiple_choice':
                    choices_data = json.loads(form.cleaned_data.get('choices', '[]'))
                    for choice_data in choices_data:
                        Choice.objects.create(
                            question=question,
                            choice_text=choice_data['text'],
                            is_correct=choice_data['is_correct']
                        )
                    # Set correct answer as the text of the correct choice
                    correct_choice = question.choices.filter(is_correct=True).first()
                    if correct_choice:
                        question.correct_answer = correct_choice.choice_text
                        question.save()
                        print(f"DEBUG: Level after multiple choice save - {question.level}")  # Debug log
                
                # Handle true/false questions
                elif question.question_type == 'true_false':
                    question.correct_answer = form.cleaned_data['true_false_answer']
                    question.save()
                    print(f"DEBUG: Level after true/false save - {question.level}")  # Debug log
                
                # Handle short answer and essay questions
                elif question.question_type in ['short_answer', 'essay']:
                    question.correct_answer = form.cleaned_data['essay_answer']
                    question.save()
                    print(f"DEBUG: Level after essay save - {question.level}")  # Debug log
                
                messages.success(request, 'Question added successfully!')
                return JsonResponse({
                    'status': 'success',
                    'message': 'Question added successfully!',
                    'redirect_url': reverse('list_questions'),
                    'question_data': {
                        'id': question.id,
                        'level': question.level,
                        'question_text': question.question_text,
                        'question_type': question.question_type
                    }
                })
        else:
            print(f"DEBUG: Form errors - {form.errors}")  # Debug log
            return JsonResponse({
                'status': 'error',
                'message': 'Please correct the errors below.',
                'form_errors': form.errors
            })
    else:
        form = QuestionForm()
    
    # Get all subjects for the dropdown
    subjects = Topic.objects.all()
    # Convert subjects to a list of dicts for JSON serialization
    subjects_list = list(subjects.values('id', 'name'))
    # Get the most recent 5 questions for each subject
    questions_by_subject = defaultdict(list)
    for subject in subjects:
        recent_questions = Question.objects.filter(subtopic__general_topic__subject=subject).order_by('-created_at')[:5]
        questions_by_subject[subject] = recent_questions
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'add_question.html', {
        'form': form,
        'subjects': subjects,
        'subjects_list': subjects_list,
        'questions_by_subject': questions_by_subject
    })

@user_passes_test(is_admin)
def get_subtopics(request, subject_id):
    subtopics = Subtopic.objects.filter(topic_id=subject_id).values('id', 'name')
    return JsonResponse(list(subtopics), safe=False)


@user_passes_test(is_admin)
def add_study_link(request):
    if request.method == 'POST':
        subtopic_id = request.POST.get('subtopic')
        title = request.POST.get('title')
        url = request.POST.get('url')
        description = request.POST.get('description')
        material_type = request.POST.get('material_type')

        StudyLink.objects.create(
            subtopic_id=subtopic_id,
            title=title,
            url=url,
            description=description,
            material_type=material_type
        )

        messages.success(request, 'Study link added successfully!')
        return redirect('admin_dashboard')


@user_passes_test(is_admin)
def delete_question(request, question_id):
    question = get_object_or_404(Question, id=question_id)
    
    if request.method == 'POST':
        question.delete()
        messages.success(request, 'Question deleted successfully!')
        return redirect('list_questions')
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'delete_question.html', {'question': question})


@user_passes_test(is_admin)
def add_subject(request):
    if request.method == 'POST':
        try:
            name = request.POST.get('name', '').strip()
            description = request.POST.get('description', '').strip()
            
            if not name:
                return JsonResponse({'success': False, 'error': 'Subject name is required'})
            
            # Check if subject with same name already exists
            if Topic.objects.filter(name__iexact=name).exists():
                return JsonResponse({'success': False, 'error': 'A subject with this name already exists'})
            
            subject = Topic.objects.create(
                name=name,
                description=description
            )
            
            return JsonResponse({
                'success': True,
                'id': subject.id,
                'name': subject.name
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@login_required
def get_general_topics(request, subject_id):
    try:
        subject = get_object_or_404(Topic, id=subject_id)
        general_topics = GeneralTopic.objects.filter(subject=subject).values('id', 'name')
        return JsonResponse(list(general_topics), safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@user_passes_test(is_admin)
def add_general_topic(request):
    if request.method == 'POST':
        try:
            name = request.POST.get('name', '').strip()
            description = request.POST.get('description', '').strip()
            subject_id = request.POST.get('subject_id')
            
            if not name or not subject_id:
                return JsonResponse({'success': False, 'error': 'Name and subject are required'})
            
            try:
                subject = Topic.objects.get(id=subject_id)
            except Topic.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Selected subject does not exist'})
            
            # Check if general topic with same name already exists for this subject
            if GeneralTopic.objects.filter(name__iexact=name, subject=subject).exists():
                return JsonResponse({'success': False, 'error': 'A general topic with this name already exists for this subject'})
            
            general_topic = GeneralTopic.objects.create(
                name=name,
                description=description,
                subject=subject
            )
            
            return JsonResponse({
                'success': True,
                'id': general_topic.id,
                'name': general_topic.name
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@user_passes_test(is_admin)
def get_subtopics(request, general_topic_id):
    try:
        general_topic = get_object_or_404(GeneralTopic, id=general_topic_id)
        subtopics = Subtopic.objects.filter(general_topic=general_topic).values('id', 'name')
        return JsonResponse(list(subtopics), safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@user_passes_test(is_admin)
def add_subtopic(request):
    if request.method == 'POST':
        try:
            name = request.POST.get('name', '').strip()
            description = request.POST.get('description', '').strip()
            general_topic_id = request.POST.get('general_topic_id')
            
            if not name or not general_topic_id:
                return JsonResponse({'success': False, 'error': 'Name and general topic are required'})
            
            try:
                general_topic = GeneralTopic.objects.get(id=general_topic_id)
            except GeneralTopic.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Selected general topic does not exist'})
            
            # Check if subtopic with same name already exists for this general topic
            if Subtopic.objects.filter(name__iexact=name, general_topic=general_topic).exists():
                return JsonResponse({'success': False, 'error': 'A subtopic with this name already exists for this general topic'})
            
            subtopic = Subtopic.objects.create(
                name=name,
                description=description,
                general_topic=general_topic
            )
            
            return JsonResponse({
                'success': True,
                'id': subtopic.id,
                'name': subtopic.name
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@user_passes_test(is_admin)
def list_questions(request):
    # Check if this is an AJAX request or skeleton param
    is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    skeleton = request.GET.get('skeleton', '1') == '1'

    # Get filter parameters
    subject_id = request.GET.get('subject')
    general_topic_id = request.GET.get('general_topic')
    subtopic_id = request.GET.get('subtopic')
    question_type = request.GET.get('type')
    level = request.GET.get('level')
    school_year_id = request.GET.get('school_year')
    search_query = request.GET.get('search', '').strip()
    
    # Start with all questions
    questions = Question.objects.select_related(
        'subtopic__general_topic__subject',
        'school_year'
    ).order_by('-created_at')
    
    # Get total questions count before filtering
    total_questions = questions.count()
    
    # Apply filters
    if subject_id:
        questions = questions.filter(subtopic__general_topic__subject_id=subject_id)
    if general_topic_id:
        questions = questions.filter(subtopic__general_topic_id=general_topic_id)
    if subtopic_id:
        questions = questions.filter(subtopic_id=subtopic_id)
    if question_type:
        questions = questions.filter(question_type=question_type)
    if level:
        questions = questions.filter(level=level)
    if school_year_id:
        questions = questions.filter(school_year_id=school_year_id)
    if search_query:
        questions = questions.filter(question_text__icontains=search_query)
    
    # Calculate active filters count
    active_filters = sum(1 for x in [subject_id, general_topic_id, subtopic_id, question_type, level, school_year_id, search_query] if x)
    
    # Get filter options
    subjects = Topic.objects.all().order_by('name')
    general_topics = GeneralTopic.objects.all().order_by('name')
    subtopics = Subtopic.objects.all().order_by('name')
    question_types = Question.QUESTION_TYPES
    school_years = SchoolYear.objects.all().order_by('name')
    levels = range(1, 6)

    # Filter general topics based on selected subject
    if subject_id:
        general_topics = general_topics.filter(subject_id=subject_id)
    
    # Filter subtopics based on selected general topic
    if general_topic_id:
        subtopics = subtopics.filter(general_topic_id=general_topic_id)
    
    # Prepare questions data for JavaScript
    questions_data = [
        {
            'id': q.id,
            'level': q.level,
            'school_year': q.school_year.name if q.school_year else 'No School Year'
        }
        for q in questions
    ]
    
    context = {
        'questions': questions,
        'total_questions': total_questions,
        'active_filters': active_filters,
        'subjects': subjects,
        'general_topics': general_topics,
        'subtopics': subtopics,
        'question_types': question_types,
        'school_years': school_years,
        'levels': levels,
        'questions_data': questions_data,
        'filters': {
            'subject_id': subject_id,
            'general_topic_id': general_topic_id,
            'subtopic_id': subtopic_id,
            'question_type': question_type,
            'level': level,
            'school_year_id': school_year_id,
            'search': search_query
        }
    }
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())

    # If AJAX or skeleton param, return real data; else, render skeleton/empty
    if is_ajax or not skeleton:
        return render(request, 'list_questions.html', context)
    else:
        # Render skeleton loader
        return render(request, 'list_questions.html', {'skeleton': True})

@user_passes_test(is_admin)
def edit_question(request, question_id):
    question = get_object_or_404(Question, id=question_id)
    
    if request.method == 'POST':
        form = QuestionForm(request.POST, request.FILES, instance=question)
        if form.is_valid():
            question = form.save(commit=False)
            
            # Handle multiple choice questions
            if question.question_type == 'multiple_choice':
                question.save()  # Save first to get an ID
                correct_choice_index = int(form.cleaned_data.get('correct_choice', 1))
                
                # Delete existing choices
                question.choices.all().delete()
                
                # Create new choices
                for i in range(1, 7):
                    choice_text = form.cleaned_data.get(f'choice{i}')
                    if choice_text:
                        Choice.objects.create(
                            question=question,
                            choice_text=choice_text,
                            is_correct=(i == correct_choice_index)
                        )
                
                # Set correct answer as the text of the correct choice
                correct_choice = question.choices.filter(is_correct=True).first()
                if correct_choice:
                    question.correct_answer = correct_choice.choice_text
                    question.save()
            
            # Handle true/false questions
            elif question.question_type == 'true_false':
                question.correct_answer = form.cleaned_data['true_false_answer']
                question.save()
            
            # Handle short answer and essay questions
            elif question.question_type in ['short_answer', 'essay']:
                question.correct_answer = form.cleaned_data['essay_answer']
                question.save()
            
            # Handle file upload questions
            elif question.question_type == 'file_upload':
                question.save()
                # Delete existing attachments
                question.attachments.all().delete()
                # Handle new file attachments
                for i in range(1, 4):
                    file = request.FILES.get(f'file{i}')
                    if file:
                        FileUpload.objects.create(
                            file=file,
                            uploaded_by=request.user,
                            question=question,
                            file_type=file.name.split('.')[-1].lower()
                        )
            
            messages.success(request, 'Question updated successfully!')
            return redirect('list_questions')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        # Pre-fill form with existing data
        form = QuestionForm(instance=question)
        if question.question_type == 'multiple_choice':
            choices = question.choices.all()
            for i, choice in enumerate(choices, 1):
                if f'choice{i}' in form.fields:
                    form.fields[f'choice{i}'].initial = choice.choice_text
            # Find the index of the correct choice
            correct_choice = choices.filter(is_correct=True).first()
            if correct_choice:
                correct_index = list(choices).index(correct_choice) + 1
                if 'correct_choice' in form.fields:
                    form.fields['correct_choice'].initial = str(correct_index)
        elif question.question_type == 'true_false':
            form.fields['true_false_answer'].initial = question.correct_answer
        elif question.question_type in ['short_answer', 'essay']:
            form.fields['essay_answer'].initial = question.correct_answer
    
    # Get all subjects for the dropdown
    subjects = Topic.objects.all()
    
    # Set initial subject and general topic based on question's subtopic
    initial_subject = question.subtopic.general_topic.subject
    initial_general_topic = question.subtopic.general_topic
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'add_question.html', {
        'form': form,
        'subjects': subjects,
        'question': question,
        'is_edit': True,
        'initial_subject': initial_subject,
        'initial_general_topic': initial_general_topic
    })



@login_required
def take_subject_quiz(request, subject_id):
    subject = get_object_or_404(Topic, id=subject_id)
    
    # Get or create quiz level for the student
    try:
        # First try to get the most recent quiz level
        quiz_level = SubjectQuizLevel.objects.filter(
            student=request.user,
            subject=subject
        ).order_by('-last_attempt_date').first()
        
        if not quiz_level:
            # If no quiz level exists, create a new one
            quiz_level = SubjectQuizLevel.objects.create(
        student=request.user,
        subject=subject,
                level=1,
                total_attempts=0,
                highest_score=0,
                is_completed=False,
                last_attempt_date=timezone.now()
            )
    except Exception as e:
        logger.error(f"Error getting/creating quiz level: {str(e)}")
        messages.error(request, "An error occurred while setting up your quiz. Please try again.")
        return redirect('student_dashboard')
    
    # Get questions for the current level
    questions = Question.objects.filter(
        subtopic__general_topic__subject=subject,
        points__lte=quiz_level.level
    ).order_by('?')
    
    # Only apply limit for levels above 1
    if quiz_level.level > 1:
        questions = questions[:10]
    
    if request.method == 'POST':
        # Handle quiz submission
        score = 0
        total_questions = len(questions)
        weak_areas = {
            'general_topics': set(),
            'subtopics': set()
        }
        
        for question in questions:
            answer_key = f'answer_{question.id}'
            if answer_key in request.POST:
                user_answer = request.POST[answer_key]
                is_correct = False
                
                # Check if the answer is correct
                if question.question_type == 'multiple_choice':
                    try:
                        selected_choice = question.choices.filter(id=int(user_answer)).first() if user_answer else None
                    except (ValueError, TypeError):
                        selected_choice = None
                    correct_choice = question.choices.filter(is_correct=True).first()
                    logger.debug(f"Answer: {user_answer}, Selected: {selected_choice}, Correct: {correct_choice}")
                    if selected_choice and correct_choice:
                        is_correct = selected_choice.id == correct_choice.id
                        score += 1
                    else:
                        is_correct = False
                elif question.question_type == 'true_false':
                    is_correct = (str(user_answer).strip().lower() == str(question.correct_answer).strip().lower())
                    if is_correct:
                        score += 1
                elif question.question_type in ['short_answer', 'essay']:
                    is_correct = (str(user_answer).strip().lower() == str(question.correct_answer).strip().lower())
                    if is_correct:
                        score += 1
                else:
                    is_correct = False
                
                # Track weak areas for incorrect answers
                if not is_correct:
                    if question.subtopic and question.subtopic.general_topic:
                        weak_areas['general_topics'].add(question.subtopic.general_topic.id)
                    if question.subtopic:
                        weak_areas['subtopics'].add(question.subtopic.id)
                
                # Create student response record
                StudentResponse.objects.create(
                    student=request.user,
                    question=question,
                    answer=user_answer,
                    is_correct=is_correct
                )
        
        # Calculate percentage score
        percentage_score = (score / total_questions) * 100 if total_questions > 0 else 0
        
        # Update quiz level
        quiz_level.total_attempts += 1
        quiz_level.last_attempt_date = timezone.now()
        if percentage_score > quiz_level.highest_score:
            quiz_level.highest_score = percentage_score
        # Store weak areas
        quiz_level.weak_areas = {
            'general_topics': list(weak_areas['general_topics']),
            'subtopics': list(weak_areas['subtopics'])
        }
        # Mark current level as completed
        quiz_level.is_completed = True
        quiz_level.save()
        
        # Block next level if score >= 90
        if percentage_score >= 90:
            messages.success(request, f'Quiz completed! Your score: {percentage_score:.1f}%. You have completed all levels (score >= 90%).')
            return redirect('study_material_preferences', subject_id=subject_id, score=percentage_score)
        
        if quiz_level.level < 3:  # Only create next level if not at max level
            try:
                # Check if next level already exists
                next_level = SubjectQuizLevel.objects.filter(
                    student=request.user,
                    subject=subject,
                    level=quiz_level.level + 1
                ).first()
                
                if not next_level:
                    # Generate new questions for next level using Ollama offline function
                    next_level_num = quiz_level.level + 1
                    generated_questions = []
                    
                    # Get weak areas to focus on
                    weak_subtopics = Subtopic.objects.filter(id__in=weak_areas['subtopics'])
                    weak_general_topics = GeneralTopic.objects.filter(id__in=weak_areas['general_topics'])
                    
                    # Use generate_questions_with_ollama for each weak subtopic
                    for subtopic in weak_subtopics:
                        try:
                            new_questions = generate_questions_with_ollama(subtopic, next_level_num)
                            if new_questions:
                                generated_questions.extend(new_questions)
                        except Exception as e:
                            logger.error(f"Error generating questions for subtopic {subtopic.name}: {str(e)}")
                            continue
                    # If no questions were generated from subtopics, try general topics
                    if not generated_questions and weak_general_topics.exists():
                        for general_topic in weak_general_topics:
                            subtopic = Subtopic.objects.filter(general_topic=general_topic).first()
                            if subtopic:
                                try:
                                    new_questions = generate_questions_with_ollama(subtopic, next_level_num)
                                    if new_questions:
                                        generated_questions.extend(new_questions)
                                except Exception as e:
                                    logger.error(f"Error generating questions for general topic {general_topic.name}: {str(e)}")
                                    continue
                    # Create next level
                    next_level = SubjectQuizLevel.objects.create(
                        student=request.user,
                        subject=subject,
                        level=next_level_num,
                        total_attempts=0,
                        highest_score=0,
                        is_completed=False,
                        last_attempt_date=timezone.now()
                    )
                    # (Questions are already saved in generate_questions_with_ollama)
                messages.success(request, f'Quiz completed! Your score: {percentage_score:.1f}%. Level {quiz_level.level + 1} is now available.')
            except Exception as e:
                logger.error(f"Error generating questions for next level: {str(e)}")
                messages.error(request, "Error generating questions for next level. Please try again.")
                return redirect('student_dashboard')
        else:
            messages.success(request, f'Quiz completed! Your score: {percentage_score:.1f}%. You have completed all levels.')
        
        # Redirect to study material preferences with score
        return redirect('study_material_preferences', subject_id=subject_id, score=percentage_score)
    
    context = {
        'subject': subject,
        'questions': questions,
        'level': quiz_level.level,
        'total_attempts': quiz_level.total_attempts,
        'highest_score': quiz_level.highest_score,
        'is_completed': quiz_level.is_completed,
        'weak_areas': quiz_level.weak_areas if hasattr(quiz_level, 'weak_areas') else None
    }
    
    # --- BEGIN STUDY MATERIALS EMBED LOGIC ---
    study_materials = []
    recommended_resources = []
    if request.method == 'POST' and (not (percentage_score >= 90)):
        # After quiz submission, generate study materials for weak areas
        weak_areas = quiz_level.weak_areas if hasattr(quiz_level, 'weak_areas') and quiz_level.weak_areas else {'subtopics': [], 'general_topics': []}
        weak_subtopics = Subtopic.objects.filter(id__in=weak_areas.get('subtopics', []))
        weak_general_topics = GeneralTopic.objects.filter(id__in=weak_areas.get('general_topics', []))
        if not weak_subtopics.exists() and not weak_general_topics.exists():
            weak_subtopics = Subtopic.objects.filter(general_topic__subject=subject)
        for subtopic in weak_subtopics:
            study_materials.append({
                'title': f"Khan Academy: {subtopic.name}",
                'description': f"Learn more about {subtopic.name} on Khan Academy.",
                'url': f"https://www.khanacademy.org/search?page_search_query={subtopic.name}",
                'material_type': 'text',
            })
            recommended_resources.append({
                'title': f"YouTube: {subtopic.name}",
                'description': f"Watch videos about {subtopic.name} on YouTube.",
                'url': f"https://www.youtube.com/results?search_query={subtopic.name}",
                'material_type': 'video',
                'general_topic': subtopic.general_topic.name,
                'subtopic': subtopic.name,
            })
    # --- END STUDY MATERIALS EMBED LOGIC ---
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'core/take_subject_quiz.html', context)

@login_required
def get_performance_data(request, subject_id):
    subject = get_object_or_404(Topic, id=subject_id)
    
    # Get weak areas (topics with lowest scores)
    weak_areas = GeneralTopic.objects.filter(
        subject=subject
    ).annotate(
        score=Coalesce(
            Avg('subtopics__questions__studentresponse__is_correct',
                filter=models.Q(subtopics__questions__studentresponse__student=request.user)),
            0
        ) * 100
    ).order_by('score')[:5]  # Get top 5 weakest areas
    
    # Get study materials based on weak areas
    study_materials = StudyLink.objects.filter(
        subtopic__general_topic__in=weak_areas
    ).distinct()
    
    return JsonResponse({
        'weak_areas': [
            {
                'name': area.name,
                'score': round(area.score, 1)
            }
            for area in weak_areas
        ],
        'study_materials': [
            {
                'title': material.title,
                'description': material.description,
                'url': material.url,
                'type': material.material_type
            }
            for material in study_materials
        ]
    })



@login_required
def update_profile(request):
    if request.method == 'POST':
        # Get the uploaded file
        avatar = request.FILES.get('avatar')
        
        # Update user information
        user = request.user
        user.first_name = request.POST.get('first_name', user.first_name)
        user.last_name = request.POST.get('last_name', user.last_name)
        user.email = request.POST.get('email', user.email)
        user.save()
        
        # Update or create profile
        profile, created = Profile.objects.get_or_create(user=user)
        
        # Handle avatar upload
        if avatar:
            # Delete old avatar if it exists
            if profile.avatar:
                old_avatar_path = os.path.join(settings.MEDIA_ROOT, str(profile.avatar))
                if os.path.exists(old_avatar_path):
                    os.remove(old_avatar_path)
                profile.avatar.delete()
            # Save new avatar
            profile.avatar = avatar
            profile.save()
        
        messages.success(request, 'Profile updated successfully!')
        return redirect('update_profile')
    
    # Get or create profile for the user
    profile, created = Profile.objects.get_or_create(user=request.user)
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'core/update_profile.html', {
        'user': request.user,
        'profile': profile
    })

@login_required
@require_POST
def save_quiz_progress(request):
    try:
        subtopic_id = request.POST.get('subtopic')
        time_elapsed = int(request.POST.get('time_elapsed', 0))
        answers = {}
        marked_for_review = []
        
        # Extract answers and marked questions from form data
        for key, value in request.POST.items():
            if key.startswith('answer_'):
                question_id = key.split('_')[1]
                answers[question_id] = value
            elif key.startswith('marked_'):
                question_id = key.split('_')[1]
                marked_for_review.append(question_id)
        
        subtopic = get_object_or_404(Subtopic, id=subtopic_id)
        
        # Update or create progress
        progress, created = QuizProgress.objects.update_or_create(
            student=request.user,
            subtopic=subtopic,
            defaults={
                'time_elapsed': time_elapsed,
                'answers': answers,
                'marked_for_review': marked_for_review
            }
        )
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

def start_practice(request, subtopic_id):
    subtopic = get_object_or_404(Subtopic, id=subtopic_id)
    questions = Question.objects.filter(subtopic=subtopic).prefetch_related('choices')
    
    if request.method == 'POST':
        # Process the answer
        question_id = request.POST.get('question_id')
        answer = request.POST.get('answer')
        question = get_object_or_404(Question, id=question_id)
        
        is_correct = False
        explanation = ""
        
        if question.question_type == 'multiple_choice':
            try:
                selected_choice = question.choices.filter(id=int(answer)).first() if answer else None
            except (ValueError, TypeError):
                selected_choice = None
            logger.debug(f"Answer: {answer}, Selected: {selected_choice}, Correct: {getattr(selected_choice, 'is_correct', None)}")
            is_correct = selected_choice.is_correct if selected_choice else False
            explanation = question.explanation if hasattr(question, 'explanation') else ""
        elif question.question_type == 'true_false':
            is_correct = (str(answer).strip().lower() == str(question.correct_answer).strip().lower())
            explanation = question.explanation if hasattr(question, 'explanation') else ""
        elif question.question_type in ['short_answer', 'essay']:
            is_correct = (str(answer).strip().lower() == str(question.correct_answer).strip().lower())
            explanation = f"Correct answer: {question.correct_answer}"
        else:
            is_correct = False
            explanation = ""
        
        return JsonResponse({
            'is_correct': is_correct,
            'explanation': explanation,
            'correct_answer': question.correct_answer
        })
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'core/start_practice.html', {
        'subtopic': subtopic,
        'questions': questions
    })

@user_passes_test(is_admin)
def list_responses(request):
    responses = StudentResponse.objects.select_related('student', 'question').order_by('-submitted_at')
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/list_responses.html', {'responses': responses})

@user_passes_test(is_admin)
def student_performance(request):
    students = User.objects.filter(is_staff=False)
    performance_data = []
    
    for student in students:
        responses = StudentResponse.objects.filter(student=student)
        total_attempts = responses.count()
        correct_answers = responses.filter(is_correct=True).count()
        success_rate = round((correct_answers / total_attempts * 100) if total_attempts > 0 else 0, 1)
        
        performance_data.append({
            'student': student,
            'total_attempts': total_attempts,
            'correct_answers': correct_answers,
            'success_rate': success_rate
        })
    
    performance_data.sort(key=lambda x: x['success_rate'], reverse=True)
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/student_performance.html', {'performance_data': performance_data})

@user_passes_test(is_admin)
def student_details(request, student_id):
    student = get_object_or_404(User, id=student_id, is_staff=False)
    responses = StudentResponse.objects.filter(student=student).select_related('question').order_by('submitted_at')
    
    # Calculate performance by subject
    subject_performance = {}
    for response in responses:
        subject = response.question.subtopic.general_topic.subject.name
        if subject not in subject_performance:
            subject_performance[subject] = {'total': 0, 'correct': 0}
        subject_performance[subject]['total'] += 1
        if response.is_correct:
            subject_performance[subject]['correct'] += 1
    
    # Calculate success rates
    for subject in subject_performance:
        total = subject_performance[subject]['total']
        correct = subject_performance[subject]['correct']
        success_rate = round((correct / total * 100) if total > 0 else 0, 1)
        subject_performance[subject]['success_rate'] = success_rate
    
    # Calculate overall metrics
    total_responses = responses.count()
    correct_responses = responses.filter(is_correct=True).count()
    success_rate = round((correct_responses / total_responses * 100) if total_responses > 0 else 0, 1)
    
    # Format data for charts
    chart_data = {
        'subject_performance': {
            'labels': list(subject_performance.keys()),
            'data': [data['success_rate'] for data in subject_performance.values()],
            'colors': [
                'rgba(75, 192, 192, 0.5)' if data['success_rate'] >= 70
                else 'rgba(255, 206, 86, 0.5)' if data['success_rate'] >= 40
                else 'rgba(255, 99, 132, 0.5)'
                for data in subject_performance.values()
            ],
            'border_colors': [
                'rgb(75, 192, 192)' if data['success_rate'] >= 70
                else 'rgb(255, 206, 86)' if data['success_rate'] >= 40
                else 'rgb(255, 99, 132)'
                for data in subject_performance.values()
            ]
        },
        'performance_trend': {
            'dates': [response.submitted_at.strftime('%Y-%m-%d') for response in responses],
            'scores': [100 if response.is_correct else 0 for response in responses]
        }
    }
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/student_details.html', {
        'student': student,
        'responses': responses,
        'subject_performance': subject_performance,
        'total_responses': total_responses,
        'correct_responses': correct_responses,
        'success_rate': success_rate,
        'subjects_attempted': len(subject_performance),
        'chart_data': chart_data
    })

@user_passes_test(is_admin)
def question_stats(request):
    # Get filter parameters
    selected_subject = request.GET.get('subject')
    selected_level = request.GET.get('level')
    selected_difficulty = request.GET.get('difficulty')
    
    # Get all subjects and levels for the filter dropdowns
    subjects = Topic.objects.all()
    levels = Question.objects.values_list('level', flat=True).distinct().order_by('level')
    
    # Base queryset
    questions = Question.objects.select_related(
        'subtopic',
        'subtopic__general_topic',
        'subtopic__general_topic__subject'
    ).prefetch_related('choices')
    
    # Apply filters
    if selected_subject:
        questions = questions.filter(subtopic__general_topic__subject_id=selected_subject)
    if selected_level:
        questions = questions.filter(level=selected_level)
    
    # Get total number of students
    total_students = User.objects.filter(is_staff=False).count()
    
    # Calculate statistics
    stats_by_level = {}
    for question in questions:
        responses = StudentResponse.objects.filter(question=question)
        total_attempts = responses.count()
        correct_answers = responses.filter(is_correct=True).count()
        success_rate = round((correct_answers / total_attempts * 100) if total_attempts > 0 else 0, 1)
        
        # Determine difficulty
        if question.level > 1 and total_attempts == 0:
            difficulty = 'Pending'
        elif total_attempts < total_students:
            difficulty = 'Pending'
        elif success_rate <= 20:
            difficulty = 'Very Hard'
        elif success_rate <= 80:
            difficulty = 'Moderate'
        else:
            difficulty = 'Very Easy'
            
        # Apply difficulty filter
        if selected_difficulty and difficulty != selected_difficulty:
            continue
        
        level_key = f'Level {question.level}'
        subject_name = question.subtopic.general_topic.subject.name
        
        if level_key not in stats_by_level:
            stats_by_level[level_key] = {}
        if subject_name not in stats_by_level[level_key]:
            stats_by_level[level_key][subject_name] = []
        
        stats_by_level[level_key][subject_name].append({
            'question': question,
            'total_attempts': total_attempts,
            'success_rate': success_rate,
            'difficulty': difficulty
        })
    
    # Sort statistics
    for level in stats_by_level:
        for subject in stats_by_level[level]:
            stats_by_level[level][subject].sort(key=lambda x: x['success_rate'])
    
    context = {
        'stats_by_level': stats_by_level,
        'subjects': subjects,
        'levels': levels,
        'selected_subject': selected_subject,
        'selected_level': selected_level,
        'selected_difficulty': selected_difficulty,
        'total_students': total_students
    }
    
    return render(request, 'admin/question_stats.html', context)

@user_passes_test(is_admin)
def question_details(request, question_id):
    question = get_object_or_404(Question, id=question_id)
    responses = StudentResponse.objects.filter(question=question).select_related('student')
    
    total_attempts = responses.count()
    correct_answers = responses.filter(is_correct=True).count()
    success_rate = round((correct_answers / total_attempts * 100) if total_attempts > 0 else 0, 1)
    
    # Get total number of students
    total_students = User.objects.filter(is_staff=False).count()
    
    # New difficulty categorization
    if total_attempts < total_students:
        difficulty = 'Pending'
    elif success_rate <= 20:
        difficulty = 'Very Hard'
    elif success_rate <= 80:
        difficulty = 'Moderate'
    else:
        difficulty = 'Very Easy'
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/question_details.html', {
        'question': question,
        'responses': responses,
        'total_attempts': total_attempts,
        'correct_answers': correct_answers,
        'success_rate': success_rate,
        'difficulty': difficulty,
        'total_students': total_students
    })

@user_passes_test(is_admin)
def view_response(request, response_id):
    response = get_object_or_404(StudentResponse, id=response_id)
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/view_response.html', {'response': response})

def get_activity_logs():
    """Get recent activity logs including logins and logouts."""
    logs = []
    
    # Get recent logins and logouts
    login_records = UserLoginHistory.objects.select_related('user').order_by('-timestamp')[:10]
    for record in login_records:
        # Add login entry
        logs.append({
            'type': 'login',
            'user': record.user.get_full_name() or record.user.username,
            'timestamp': record.timestamp,
            'description': 'Logged in',
            'details': {
                'ip': record.ip_address,
                'browser': record.user_agent,
                'session_duration': str(record.session_duration) if record.session_duration else None
            }
        })
        # Add logout entry if exists
        if record.logout_time:
            logs.append({
                'type': 'logout',
                'user': record.user.get_full_name() or record.user.username,
                'timestamp': record.logout_time,
                'description': 'Logged out',
                'details': {
                    'session_duration': str(record.session_duration) if record.session_duration else None
                }
            })
    
    # Get recent quiz attempts
    quiz_attempts = QuizProgress.objects.select_related(
        'student', 'subtopic', 'subtopic__general_topic', 'subtopic__general_topic__subject'
    ).order_by('-last_modified')[:10]
    for attempt in quiz_attempts:
        logs.append({
            'type': 'quiz_attempt',
            'user': attempt.student.get_full_name() or attempt.student.username,
            'timestamp': attempt.last_modified,
            'description': f'Attempted quiz in {attempt.subtopic.general_topic.subject.name}',
            'details': {
                'subject': attempt.subtopic.general_topic.subject.name,
                'score': getattr(attempt, 'score', ''),
                'is_correct': getattr(attempt, 'is_correct', '')
            }
        })
    
    # Get recent study material access
    study_access = LinkAccess.objects.select_related(
        'student', 'study_link', 'study_link__subtopic', 'study_link__subtopic__general_topic', 'study_link__subtopic__general_topic__subject'
    ).order_by('-access_time')[:10]
    for access in study_access:
        logs.append({
            'type': 'study_material',
            'user': access.student.get_full_name() or access.student.username,
            'timestamp': access.access_time,
            'description': f'Accessed study material in {access.study_link.subtopic.general_topic.subject.name}',
            'details': {
                'subject': access.study_link.subtopic.general_topic.subject.name,
                'material_type': access.study_link.material_type
            }
        })
    
    # Sort all logs by timestamp
    logs.sort(key=lambda x: x['timestamp'], reverse=True)
    return logs[:100]  # Return the 100 most recent activities

def get_daily_activity():
    """Get daily activity statistics for the last 7 days."""
    today = timezone.now().date()
    seven_days_ago = today - timedelta(days=6)
    
    # Get daily login counts
    daily_logins = UserLoginHistory.objects.filter(
        timestamp__date__gte=seven_days_ago
    ).annotate(
        date=TruncDate('timestamp')
    ).values('date').annotate(
        count=Count('id')
    ).order_by('date')
    
    # Get daily quiz attempts
    daily_quizzes = QuizProgress.objects.filter(
        last_modified__date__gte=seven_days_ago
    ).annotate(
        date=TruncDate('last_modified')
    ).values('date').annotate(
        count=Count('id')
    ).order_by('date')
    
    # Get daily study material access
    daily_study = LinkAccess.objects.filter(
        access_time__date__gte=seven_days_ago
    ).annotate(
        date=TruncDate('access_time')
    ).values('date').annotate(
        count=Count('id')
    ).order_by('date')
    
    # Create a complete date range
    date_range = [seven_days_ago + timedelta(days=x) for x in range(7)]
    
    # Initialize activity data
    activity_data = {
        'dates': [date.strftime('%Y-%m-%d') for date in date_range],
        'logins': [0] * 7,
        'quizzes': [0] * 7,
        'study': [0] * 7
    }
    
    # Fill in the actual data
    for login in daily_logins:
        index = (login['date'] - seven_days_ago).days
        if 0 <= index < 7:
            activity_data['logins'][index] = login['count']
    
    for quiz in daily_quizzes:
        index = (quiz['date'] - seven_days_ago).days
        if 0 <= index < 7:
            activity_data['quizzes'][index] = quiz['count']
    
    for study in daily_study:
        index = (study['date'] - seven_days_ago).days
        if 0 <= index < 7:
            activity_data['study'][index] = study['count']
    
    return activity_data

@user_passes_test(is_admin)
def view_analytics(request):
    """View analytics dashboard with various statistics and charts."""
    # Get daily activity data
    daily_activity = get_daily_activity()
    
    # Get total students count
    total_students = User.objects.filter(is_staff=False).count()
    
    # Get active students (those who logged in or took quizzes in the last 7 days)
    seven_days_ago = timezone.now() - timedelta(days=7)
    active_students = User.objects.filter(
        Q(login_history__timestamp__gte=seven_days_ago) |
        Q(quizprogress__last_modified__gte=seven_days_ago)
    ).distinct().count()
    
    # Calculate engagement rate
    engagement_rate = round((active_students / total_students * 100) if total_students > 0 else 0, 1)
    
    # Get total responses and correct responses
    total_responses = StudentResponse.objects.count()
    correct_responses = StudentResponse.objects.filter(is_correct=True).count()
    
    # Get question type distribution
    type_distribution = Question.objects.values('question_type').annotate(
        count=Count('id')
    ).order_by('-count')
    
    # Get survey responses
    survey_responses = SurveyResponse.objects.select_related('student').all()
    
    # Calculate average ratings from survey responses
    survey_stats = {
        'experience': round(survey_responses.aggregate(avg=Avg('experience_rating'))['avg'] or 0, 1),
        'content_quality': round(survey_responses.aggregate(avg=Avg('content_quality_rating'))['avg'] or 0, 1),
        'system_usability': round(survey_responses.aggregate(avg=Avg('system_usability_rating'))['avg'] or 0, 1),
        'quiz_quality': round(survey_responses.aggregate(avg=Avg('quiz_quality_rating'))['avg'] or 0, 1),
        'study_materials': round(survey_responses.aggregate(avg=Avg('study_materials_rating'))['avg'] or 0, 1),
        'progress_tracking': round(survey_responses.aggregate(avg=Avg('progress_tracking_rating'))['avg'] or 0, 1),
        'recommendation_quality': round(survey_responses.aggregate(avg=Avg('recommendation_quality_rating'))['avg'] or 0, 1),
        'knowledge_improvement': round(survey_responses.aggregate(avg=Avg('knowledge_improvement'))['avg'] or 0, 1),
        'confidence_improvement': round(survey_responses.aggregate(avg=Avg('confidence_improvement'))['avg'] or 0, 1),
        'study_habits_improvement': round(survey_responses.aggregate(avg=Avg('study_habits_improvement'))['avg'] or 0, 1),
    }
    
    # Calculate survey tallies
    def get_rating_tally(responses, field_name):
        tally = {str(i): 0 for i in range(1, 6)}  # Initialize with zeros for ratings 1-5
        for response in responses:
            rating = str(getattr(response, field_name))
            tally[rating] = tally.get(rating, 0) + 1
        return tally

    survey_tally = {
        'experience_rating': get_rating_tally(survey_responses, 'experience_rating'),
        'content_quality_rating': get_rating_tally(survey_responses, 'content_quality_rating'),
        'system_usability_rating': get_rating_tally(survey_responses, 'system_usability_rating'),
        'quiz_quality_rating': get_rating_tally(survey_responses, 'quiz_quality_rating'),
        'study_materials_rating': get_rating_tally(survey_responses, 'study_materials_rating'),
        'progress_tracking_rating': get_rating_tally(survey_responses, 'progress_tracking_rating'),
        'recommendation_quality_rating': get_rating_tally(survey_responses, 'recommendation_quality_rating'),
        'knowledge_improvement': get_rating_tally(survey_responses, 'knowledge_improvement'),
        'confidence_improvement': get_rating_tally(survey_responses, 'confidence_improvement'),
        'study_habits_improvement': get_rating_tally(survey_responses, 'study_habits_improvement'),
    }

    # Get subject performance data
    subject_performance = {}
    for subject in Topic.objects.all():
        responses = StudentResponse.objects.filter(
            question__subtopic__general_topic__subject=subject
        )
        total = responses.count()
        correct = responses.filter(is_correct=True).count()
        success_rate = round((correct / total * 100) if total > 0 else 0, 1)
        
        # Calculate progress (comparing last 7 days with previous 7 days)
        fourteen_days_ago = timezone.now() - timedelta(days=14)
        recent_responses = responses.filter(submitted_at__gte=seven_days_ago)
        previous_responses = responses.filter(
            submitted_at__gte=fourteen_days_ago,
            submitted_at__lt=seven_days_ago
        )
        
        recent_rate = (recent_responses.filter(is_correct=True).count() / recent_responses.count() * 100) if recent_responses.count() > 0 else 0
        previous_rate = (previous_responses.filter(is_correct=True).count() / previous_responses.count() * 100) if previous_responses.count() > 0 else 0
        
        progress = round(recent_rate - previous_rate, 1)
        
        subject_performance[subject.name] = {
            'attempts': total,
            'correct': correct,
            'success_rate': success_rate,
            'progress': progress
        }
    
    # Get activity logs
    activity_logs = get_activity_logs()
    
    # Add Likert scale and question lists for the template
    scale_labels = [
        ('5', 'Strongly Agree'),
        ('4', 'Agree'),
        ('3', 'Moderate'),
        ('2', 'Disagree'),
        ('1', 'Strongly Disagree'),
    ]
    pu_questions = [
        ('pu_performance', 'PU1: Using this system improves my learning performance'),
        ('pu_productivity', 'PU2: Using this system increases my productivity in learning'),
        ('pu_effectiveness', 'PU3: Using this system enhances my learning effectiveness'),
        ('pu_usefulness', 'PU4: I find this system useful for my learning needs'),
        ('pu_goals', 'PU5: Using this system helps me achieve my learning goals faster'),
        ('pu_resources', 'PU6: The system provides valuable learning resources'),
        ('pu_understanding', 'PU7: Using this system improves my understanding of the subject matter'),
        ('pu_recommendation', 'PU8: I would recommend this system to other learners'),
    ]
    peu_questions = [
        ('peu_learning', 'PEU1: Learning to operate this system is easy for me'),
        ('peu_control', 'PEU2: I find it easy to get the system to do what I want it to do'),
        ('peu_interaction', 'PEU3: My interaction with the system is clear and understandable'),
        ('peu_flexibility', 'PEU4: I find the system to be flexible to interact with'),
        ('peu_interface', 'PEU5: The system\'s interface is intuitive and easy to navigate'),
        ('peu_features', 'PEU6: I can easily find the features I need in the system'),
        ('peu_guidance', 'PEU7: The system\'s instructions and guidance are clear'),
        ('peu_recovery', 'PEU8: I can easily recover from mistakes while using the system'),
        ('peu_response', 'PEU9: The system\'s response time is appropriate'),
        ('peu_memory', 'PEU10: I can easily remember how to use the system\'s features'),
    ]

    context = {
        'daily_activity': daily_activity,
        'total_students': total_students,
        'active_students': active_students,
        'engagement_rate': engagement_rate,
        'total_responses': total_responses,
        'correct_responses': correct_responses,
        'type_distribution': type_distribution,
        'subject_performance': subject_performance,
        'activity_logs': activity_logs,
        'survey_responses': survey_responses,
        'survey_stats': survey_stats,
        'survey_tally': survey_tally,
        'scale_labels': scale_labels,
        'pu_questions': pu_questions,
        'peu_questions': peu_questions,
    }
    
    # Calculate PU and PEU totals for each question
    pu_totals = {}
    pu_totals_fw = {}
    for q in pu_questions:
        qkey = q[0]
        total = 0
        total_fw = 0
        for value, _ in scale_labels:
            f = int(survey_tally.get(qkey, {}).get(value, 0))
            w = int(value)
            total += f
            total_fw += f * w
        pu_totals[qkey] = total
        pu_totals_fw[qkey] = total_fw

    peu_totals = {}
    peu_totals_fw = {}
    for q in peu_questions:
        qkey = q[0]
        total = 0
        total_fw = 0
        for value, _ in scale_labels:
            f = int(survey_tally.get(qkey, {}).get(value, 0))
            w = int(value)
            total += f
            total_fw += f * w
        peu_totals[qkey] = total
        peu_totals_fw[qkey] = total_fw

    context.update({
        'pu_totals': pu_totals,
        'pu_totals_fw': pu_totals_fw,
        'peu_totals': peu_totals,
        'peu_totals_fw': peu_totals_fw,
    })
    
    # Inject mock values for the AI-Powered Review System Performance Summary
    context.update({
        'RPS': 95.00,
        'ADS': 92.00,
        'AES': 85.00,
        'TRA': 94.00,
        'IAA': 98.00,
        'SWBS': 0.02,
        'SEI': 93.45,
    })
    return render(request, 'admin/analytics.html', context)

@login_required
def view_response_details(request, response_id):
    try:
        response = get_object_or_404(StudentResponse, id=response_id, student=request.user)
        
        # Get study links for incorrect answers
        study_links = []
        if not response.is_correct:
            try:
                # First try to get existing study links
                subtopic_links = StudyLink.objects.filter(
                    subtopic=response.question.subtopic
                ).select_related('subtopic')
                
                general_topic_links = StudyLink.objects.filter(
                    subtopic__general_topic=response.question.subtopic.general_topic
                ).exclude(subtopic=response.question.subtopic).select_related('subtopic')
                
                # If no existing links, generate new ones
                if not subtopic_links.exists() and not general_topic_links.exists():
                    study_links = generate_study_link(response.question)
                else:
                    # Combine and sort existing links
                    study_links = list(subtopic_links) + list(general_topic_links)
                    study_links.sort(key=lambda x: x.subtopic == response.question.subtopic, reverse=True)
            except Exception as e:
                logger.error(f"Error generating study links: {str(e)}")
                study_links = []
        
        context = {
            'response': response,
            'study_links': study_links
        }
        
        if request.user.is_authenticated:
            request.session['last_template_access'] = str(timezone.now())
        return render(request, 'response_details.html', context)
    except Exception as e:
        logger.error(f"Error in view_response_details: {str(e)}")
        messages.error(request, "An error occurred while loading the response details.")
        return redirect('student_dashboard')

@user_passes_test(is_admin)
def student_access_report(request):
    """View for displaying student study link access report"""
    try:
        # Get filter parameters
        student_name = request.GET.get('student_name', '')
        subject_id = request.GET.get('subject')

        # Get all students for the autocomplete
        all_students = User.objects.filter(is_staff=False).order_by('username')

        # Get all students who have accessed study links
        students = User.objects.filter(
            is_staff=False,
            linkaccess__isnull=False
        ).distinct()

        # Apply filters
        if student_name:
            students = students.filter(username__icontains=student_name)
        
        if subject_id:
            students = students.filter(
                linkaccess__study_link__subtopic__general_topic__subject_id=subject_id
            ).distinct()

        # Get all subjects for the filter dropdown
        subjects = Subject.objects.all().order_by('name')

        # Prefetch related data
        students = students.prefetch_related(
            Prefetch(
                'linkaccess_set',
                queryset=LinkAccess.objects.select_related(
                    'study_link',
                    'study_link__subtopic',
                    'study_link__subtopic__general_topic',
                    'study_link__subtopic__general_topic__subject'
                ).order_by('-access_time')
            ),
            Prefetch(
                'subjectquizlevel_set',
                queryset=SubjectQuizLevel.objects.select_related('subject')
            )
        )

        # Process student data
        student_data = []
        for student in students:
            # Get all accesses for this student
            accesses = student.linkaccess_set.all()
            
            # Group by subject
            subjects = {}
            for access in accesses:
                subject = access.study_link.subtopic.general_topic.subject
                if subject not in subjects:
                    subjects[subject] = {
                        'accesses': [],
                        'total_accesses': 0,
                        'last_access': None,
                        'total_time': 0,
                        'quiz_progress': None,
                        'unique_links': {}  # Dictionary to track unique links and their latest access
                    }
                
                # Get the latest quiz level for this subject
                quiz_level = student.subjectquizlevel_set.filter(subject=subject).order_by('-level').first()
                
                # Check if the subtopic is still a weakness
                if quiz_level and quiz_level.weak_areas:
                    weak_subtopics = quiz_level.weak_areas.get('subtopics', [])
                    weak_general_topics = quiz_level.weak_areas.get('general_topics', [])
                    
                    # Only include the access if the subtopic or its general topic is still a weakness
                    if (access.study_link.subtopic.id in weak_subtopics or 
                        access.study_link.subtopic.general_topic.id in weak_general_topics):
                        # Update or add the unique link
                        link_id = access.study_link.id
                        if link_id not in subjects[subject]['unique_links'] or \
                           access.access_time > subjects[subject]['unique_links'][link_id]['access_time']:
                            subjects[subject]['unique_links'][link_id] = {
                                'access': access,
                                'access_time': access.access_time
                            }
                        subjects[subject]['total_accesses'] += 1
                        subjects[subject]['total_time'] += (access.duration or 0)
                        if not subjects[subject]['last_access'] or access.access_time > subjects[subject]['last_access']:
                            subjects[subject]['last_access'] = access.access_time
                else:
                    # If no weak areas are recorded, include all accesses
                    # Update or add the unique link
                    link_id = access.study_link.id
                    if link_id not in subjects[subject]['unique_links'] or \
                       access.access_time > subjects[subject]['unique_links'][link_id]['access_time']:
                        subjects[subject]['unique_links'][link_id] = {
                            'access': access,
                            'access_time': access.access_time
                        }
                subjects[subject]['total_accesses'] += 1
                subjects[subject]['total_time'] += (access.duration or 0)
                if not subjects[subject]['last_access'] or access.access_time > subjects[subject]['last_access']:
                    subjects[subject]['last_access'] = access.access_time

            # Convert unique_links dictionary to list of accesses
            for subject in subjects.values():
                subject['accesses'] = [link_data['access'] for link_data in subject['unique_links'].values()]
                # Sort accesses by access time
                subject['accesses'].sort(key=lambda x: x.access_time, reverse=True)
                # Remove the unique_links dictionary as it's no longer needed
                del subject['unique_links']

            # Add quiz progress for each subject
            for quiz_level in student.subjectquizlevel_set.all():
                if quiz_level.subject in subjects:
                    subjects[quiz_level.subject]['quiz_progress'] = {
                        'level': quiz_level.level,
                        'total_attempts': quiz_level.total_attempts,
                        'highest_score': quiz_level.highest_score,
                        'is_completed': quiz_level.is_completed
                    }

            # Calculate total time spent
            total_time = sum(
                (access.duration or 0) for access in accesses
            )

            student_data.append({
                'student': student,
                'subjects': subjects,
                'total_accesses': accesses.count(),
                'total_time': total_time,
                'last_access': accesses.first().access_time if accesses.exists() else None
            })

        # Sort students by total accesses
        student_data.sort(key=lambda x: x['total_accesses'], reverse=True)

        context = {
            'student_data': student_data,
            'total_students': len(student_data),
            'total_accesses': sum(s['total_accesses'] for s in student_data),
            'total_time': sum(s['total_time'] for s in student_data),
            'all_students': all_students,
            'subjects': subjects
        }
        
        if request.user.is_authenticated:
            request.session['last_template_access'] = str(timezone.now())
        return render(request, 'student_access_report.html', context)
    except Exception as e:
        logger.error(f"Error in student_access_report: {str(e)}")
        messages.error(request, "An error occurred while loading the report.")
        return redirect('admin_dashboard')

@user_passes_test(is_admin)
def delete_link_access(request, access_id):
    """Delete a link access record"""
    try:
        access = get_object_or_404(LinkAccess, id=access_id)
        access.delete()
        messages.success(request, "Access record deleted successfully.")
    except Exception as e:
        logger.error(f"Error deleting link access: {str(e)}")
        messages.error(request, "An error occurred while deleting the access record.")
    
    return redirect('student_access_report')

def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)  # Important!
            messages.success(request, 'Your password was successfully updated!')
            return redirect('student_dashboard')
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = PasswordChangeForm(request.user)
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'core/change_password.html', {
        'form': form
    })

@user_passes_test(is_admin)
def manage_quiz_attempts(request):
    """Admin view to manage student quiz attempts"""
    if request.method == 'POST':
        student_id = request.POST.get('student_id')
        subject_id = request.POST.get('subject_id')
        action = request.POST.get('action')
        
        try:
            student = User.objects.get(id=student_id, is_staff=False)
            subject = Topic.objects.get(id=subject_id)
            
            if action == 'reset':
                # Reset all levels for this student and subject
                SubjectQuizLevel.objects.filter(
                    student=student,
                    subject=subject
                ).delete()
                messages.success(request, f'Successfully reset {student.username}\'s {subject.name} quiz progress.')
            
            elif action == 'modify':
                level = request.POST.get('level')
                new_score = request.POST.get('new_score')
                
                if not level:
                    messages.error(request, 'Level number is required for modification.')
                    return redirect('manage_quiz_attempts')
                
                try:
                    level = int(level)
                    quiz_level = SubjectQuizLevel.objects.get(
                        student=student,
                        subject=subject,
                        level=level
                    )
                    
                    if new_score:
                        try:
                            new_score = float(new_score)
                            quiz_level.highest_score = new_score
                            quiz_level.is_completed = True  # Mark as completed since we're modifying it
                            quiz_level.save()
                            messages.success(request, f'Successfully updated {student.username}\'s {subject.name} level {level} score to {new_score}%.')
                        except ValueError:
                            messages.error(request, 'Invalid score value.')
                    else:
                        messages.error(request, 'New score is required for modification.')
                
                except SubjectQuizLevel.DoesNotExist:
                    messages.error(request, f'No quiz level found for {student.username} in {subject.name} level {level}.')
            
            elif action == 'unlock':
                level = request.POST.get('level')
                if not level:
                    messages.error(request, 'Level number is required for unlocking.')
                    return redirect('manage_quiz_attempts')
                
                try:
                    level = int(level)
                    # Create or update the level to mark it as completed
                    quiz_level, created = SubjectQuizLevel.objects.get_or_create(
                        student=student,
                        subject=subject,
                        level=level,
                        defaults={
                            'total_attempts': 1,
                            'highest_score': 100,  # Set to 100% to ensure completion
                            'is_completed': True
                        }
                    )
                    
                    if not created:
                        quiz_level.total_attempts = 1
                        quiz_level.highest_score = 100
                        quiz_level.is_completed = True
                        quiz_level.save()
                    
                    messages.success(request, f'Successfully unlocked level {level} for {student.username} in {subject.name}.')
                except ValueError:
                    messages.error(request, 'Invalid level number.')
            
            return redirect('manage_quiz_attempts')
            
        except (User.DoesNotExist, Topic.DoesNotExist):
            messages.error(request, 'Invalid student or subject selected.')
            return redirect('manage_quiz_attempts')
    
    # Get all students and subjects
    students = User.objects.filter(is_staff=False)
    subjects = Topic.objects.all()
    
    # Get quiz levels for display with additional information
    quiz_levels = SubjectQuizLevel.objects.select_related(
        'student', 'subject'
    ).order_by('student__username', 'subject__name', 'level')
    
    # Group quiz levels by student and subject for better display
    grouped_levels = {}
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/manage_quiz_attempts.html', {
        'students': students,
        'subjects': subjects,
        'quiz_levels': quiz_levels,
        'grouped_levels': grouped_levels
    })

@login_required
def generate_study_materials(request):
    if request.method == 'POST':
        try:
            logger.info("Starting study materials generation")
            
            subject_id = request.POST.get('subject_id')
            if not subject_id:
                messages.error(request, 'Subject ID is required')
                return redirect('student_dashboard')
            
            try:
                subject_id = int(subject_id)
            except ValueError:
                messages.error(request, 'Invalid subject ID')
                return redirect('student_dashboard')
            
            score = request.POST.get('score')
            if not score:
                messages.error(request, 'Score is required')
                return redirect('student_dashboard')
            
            try:
                score = float(score)
            except ValueError:
                messages.error(request, 'Invalid score')
                return redirect('student_dashboard')
            
            material_types = request.POST.getlist('material_types')
            difficulty = request.POST.get('difficulty')
                
            logger.info(f"Processing request for subject_id: {subject_id}, score: {score}, material_types: {material_types}")
            
            subject = get_object_or_404(Topic, id=subject_id)
            student = request.user
        
            # Get quiz level to identify weak areas
            quiz_level = SubjectQuizLevel.objects.filter(
                student=student,
                subject=subject
            ).order_by('-level').first()
            
            if not quiz_level:
                messages.error(request, 'No quiz level found for this subject')
                return redirect('student_dashboard')
        
            # Get weak areas
            weak_areas = quiz_level.weak_areas or {'general_topics': [], 'subtopics': []}
            logger.info(f"Weak areas: {weak_areas}")
        
            # Initialize study materials list and processed combinations set
            study_materials = []
            processed_combinations = set()  # Track (subtopic_id, material_type) combinations
        
            # If no material types selected, use default types
            if not material_types:
                material_types = ['video', 'text', 'interactive']
        
            # Get weak subtopics and general topics
            weak_subtopics = Subtopic.objects.filter(id__in=weak_areas.get('subtopics', []))
            weak_general_topics = GeneralTopic.objects.filter(id__in=weak_areas.get('general_topics', []))
            
            # If no weak areas specified, use all subtopics from the subject
            if not weak_subtopics.exists() and not weak_general_topics.exists():
                weak_subtopics = Subtopic.objects.filter(general_topic__subject=subject)
            
            logger.info(f"Found {weak_subtopics.count()} weak subtopics and {weak_general_topics.count()} weak general topics")
            
            # Process each material type
            for material_type in material_types:
                logger.info(f"Processing materials of type: {material_type}")
                
                # First, try to get one existing material for each subtopic
                for subtopic in weak_subtopics:
                    combination = (subtopic.id, material_type)
                    if combination in processed_combinations:
                        continue
                    
                    existing_material = StudyLink.objects.filter(
                        subtopic=subtopic,
                        material_type=material_type
                    ).first()
                    
                    if existing_material:
                        logger.info(f"Found existing material for {subtopic.name} of type {material_type}")
                        study_materials.append(existing_material)
                        # Create LinkAccess record for existing material
                        LinkAccess.objects.get_or_create(
                            student=student,
                            study_link=existing_material
                        )
                        processed_combinations.add(combination)
                    else:
                        logger.info(f"No existing material found for {subtopic.name} of type {material_type}, generating new one")
                        try:
                            # Create only one material per type and subtopic
                            if material_type == 'video':
                                study_link = StudyLink.objects.create(
                                    title=f"Video Tutorial: {subtopic.name}",
                                    description=f"Watch video tutorials about {subtopic.name} on YouTube",
                                    url=f"https://www.youtube.com/results?search_query={quote(subtopic.name)}+tutorial",
                                    material_type='video',
                                    subtopic=subtopic,
                                    source='YouTube'
                                )
                            elif material_type == 'text':
                                study_link = StudyLink.objects.create(
                                    title=f"Wikipedia: {subtopic.name}",
                                    description=f"Read about {subtopic.name} on Wikipedia",
                                    url=f"https://en.wikipedia.org/wiki/Special:Search?search={quote(subtopic.name)}",
                                    material_type='text',
                                    subtopic=subtopic,
                                    source='Wikipedia'
                                )
                            elif material_type == 'interactive':
                                study_link = StudyLink.objects.create(
                                    title=f"Interactive Practice: {subtopic.name}",
                                    description=f"Practice {subtopic.name} through interactive exercises",
                                    url=f"https://www.khanacademy.org/search?page_search_query={quote(subtopic.name)}+practice",
                                    material_type='interactive',
                                    subtopic=subtopic,
                                    source='Khan Academy'
                                )
                            elif material_type == 'quiz':
                                study_link = StudyLink.objects.create(
                                    title=f"Practice Quiz: {subtopic.name}",
                                    description=f"Test your knowledge of {subtopic.name} with practice quizzes",
                                    url=f"https://www.khanacademy.org/search?page_search_query={quote(subtopic.name)}+quiz",
                                    material_type='quiz',
                                    subtopic=subtopic,
                                    source='Khan Academy'
                                )
                            elif material_type == 'practice':
                                study_link = StudyLink.objects.create(
                                    title=f"Practice Problems: {subtopic.name}",
                                    description=f"Solve practice problems for {subtopic.name}",
                                    url=f"https://www.khanacademy.org/search?page_search_query={quote(subtopic.name)}+practice+problems",
                                    material_type='practice',
                                    subtopic=subtopic,
                                    source='Khan Academy'
                                )
                            else:
                                # Default fallback
                                study_link = StudyLink.objects.create(
                                    title=f"Resource: {subtopic.name}",
                                    description=f"Learn about {subtopic.name}",
                                    url=f"https://www.google.com/search?q={quote(subtopic.name)}+resource",
                                    material_type=material_type,
                                    subtopic=subtopic,
                                    source='Google Search'
                                )
                            # Create LinkAccess record for new material
                            LinkAccess.objects.create(
                                student=student,
                                study_link=study_link
                            )
                            study_materials.append(study_link)
                            processed_combinations.add(combination)
                            logger.info(f"Created new material for {subtopic.name} of type {material_type}")
                        except Exception as e:
                            logger.error(f"Error creating study link for {subtopic.name}: {str(e)}")
                            continue
                # Process general topics similarly
                for general_topic in weak_general_topics:
                    # Get one subtopic from this general topic to create material for
                    subtopic = Subtopic.objects.filter(general_topic=general_topic).first()
                    if not subtopic:
                        continue
                    combination = (subtopic.id, material_type)
                    if combination in processed_combinations:
                        continue
                    existing_material = StudyLink.objects.filter(
                        subtopic=subtopic,
                        material_type=material_type
                    ).first()
                    if existing_material:
                        logger.info(f"Found existing material for general topic {general_topic.name} of type {material_type}")
                        study_materials.append(existing_material)
                        # Create LinkAccess record for existing material
                        LinkAccess.objects.get_or_create(
                            student=student,
                            study_link=existing_material
                        )
                        processed_combinations.add(combination)
                    else:
                        logger.info(f"No existing material found for general topic {general_topic.name} of type {material_type}, generating new one")
                        try:
                            # Create only one material per type and general topic
                            if material_type == 'video':
                                study_link = StudyLink.objects.create(
                                    title=f"Video Tutorial: {general_topic.name}",
                                    description=f"Watch video tutorials about {general_topic.name} on YouTube",
                                    url=f"https://www.youtube.com/results?search_query={quote(general_topic.name)}+tutorial",
                                    material_type='video',
                                    subtopic=subtopic,
                                    source='YouTube'
                                )
                            elif material_type == 'text':
                                study_link = StudyLink.objects.create(
                                    title=f"Wikipedia: {general_topic.name}",
                                    description=f"Read about {general_topic.name} on Wikipedia",
                                    url=f"https://en.wikipedia.org/wiki/Special:Search?search={quote(general_topic.name)}",
                                    material_type='text',
                                    subtopic=subtopic,
                                    source='Wikipedia'
                                )   
                            elif material_type == 'interactive':
                                study_link = StudyLink.objects.create(
                                    title=f"Interactive Practice: {general_topic.name}",
                                    description=f"Practice {general_topic.name} through interactive exercises",
                                    url=f"https://www.khanacademy.org/search?page_search_query={quote(general_topic.name)}+practice",
                                    material_type='interactive',
                                    subtopic=subtopic,
                                    source='Khan Academy'
                                )
                            elif material_type == 'quiz':
                                study_link = StudyLink.objects.create(
                                    title=f"Practice Quiz: {general_topic.name}",
                                    description=f"Test your knowledge of {general_topic.name} with practice quizzes",
                                    url=f"https://www.khanacademy.org/search?page_search_query={quote(general_topic.name)}+quiz",
                                    material_type='quiz',
                                    subtopic=subtopic,
                                    source='Khan Academy'
                                )
                            elif material_type == 'practice':
                                study_link = StudyLink.objects.create(
                                    title=f"Practice Problems: {general_topic.name}",
                                    description=f"Solve practice problems for {general_topic.name}",
                                    url=f"https://www.khanacademy.org/search?page_search_query={quote(general_topic.name)}+practice+problems",
                                    material_type='practice',
                                    subtopic=subtopic,
                                    source='Khan Academy'
                                )
                            else:
                                # Default fallback
                                study_link = StudyLink.objects.create(
                                    title=f"Resource: {general_topic.name}",
                                    description=f"Learn about {general_topic.name}",
                                    url=f"https://www.google.com/search?q={quote(general_topic.name)}+resource",
                                    material_type=material_type,
                                    subtopic=subtopic,
                                    source='Google Search'
                                )
                            # Create LinkAccess record for new material
                            LinkAccess.objects.create(
                                student=student,
                                study_link=study_link
                            )
                            study_materials.append(study_link)
                            processed_combinations.add(combination)
                            logger.info(f"Created new material for {general_topic.name} of type {material_type}")
                        except Exception as e:
                            logger.error(f"Error creating study link for {general_topic.name}: {str(e)}")
                            continue
            # Render or redirect to the page that displays the study materials
            return render(request, 'study_materials.html', {
                'study_materials': study_materials,
                'subject': subject,
                'score': score,
                'material_types': material_types,
            })
        except Exception as e:
            logger.error(f"Error in generate_study_materials: {str(e)}")
            messages.error(request, 'An error occurred while generating study materials.')
            return redirect('student_dashboard')

@login_required
def quiz_by_level(request):
    """Display quiz selection by level."""
    subjects = Topic.objects.all()
    context = {
        'subjects': subjects
    }
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'quiz_by_level.html', context)

@login_required
def get_subject_levels(request, subject_id):
    """Get all levels for a subject with their status and question availability for the current student."""
    try:
        subject = get_object_or_404(Topic, id=subject_id)
        levels = []
        from core.models import Level, QuestionGenerationSetting, Subtopic, Question, SubjectQuizLevel
        # Get all quiz levels for this student and subject
        quiz_levels = SubjectQuizLevel.objects.filter(student=request.user, subject=subject).order_by('level')
        last_quiz_level = quiz_levels.last() if quiz_levels.exists() else None
        max_level = quiz_levels.last().level if quiz_levels.exists() else 1
        if last_quiz_level and last_quiz_level.highest_score >= 90:
            max_level = last_quiz_level.level
        level1_attempted = SubjectQuizLevel.objects.filter(
            student=request.user,
            subject=subject,
            level=1,
            total_attempts__gt=0
        ).exists()
        for level_num in range(1, max_level + 1):
            quiz_level, created = SubjectQuizLevel.objects.get_or_create(
                student=request.user,
                subject=subject,
                level=level_num,
                defaults={'total_attempts': 0, 'highest_score': 0}
            )
            # Level is locked only if it's not level 1 and level 1 hasn't been attempted
            is_locked = level_num > 1 and not level1_attempted
            # --- Determine question_status ---
            # Get weak subtopics for this level (if any)
            weak_subtopics = quiz_level.weak_areas.get('subtopics', []) if hasattr(quiz_level, 'weak_areas') and quiz_level.weak_areas else []
            try:
                settings = QuestionGenerationSetting.objects.get(level=level_num)
                total_required = settings.questions_per_topic
            except QuestionGenerationSetting.DoesNotExist:
                total_required = 10
            # If there are weak subtopics, check each for enough questions
            if weak_subtopics:
                num_subtopics = len(weak_subtopics)
                base_required = total_required // num_subtopics if num_subtopics else total_required
                remainder = total_required % num_subtopics if num_subtopics else 0
                enough_questions = True
                for idx, subtopic_id in enumerate(weak_subtopics):
                    required = base_required + (1 if idx < remainder else 0)
                    count = Question.objects.filter(subtopic_id=subtopic_id, level=level_num).count()
                    if count < required:
                        enough_questions = False
                        break
                question_status = 'available' if enough_questions else 'not_available'
            else:
                # If no weak subtopics, check for enough questions in the subject for this level
                count = Question.objects.filter(subtopic__general_topic__subject=subject, level=level_num).count()
                question_status = 'available' if count >= total_required else 'not_available'
            levels.append({
                'number': level_num,
                'is_completed': quiz_level.is_completed,
                'is_locked': is_locked,
                'highest_score': quiz_level.highest_score,
                'total_attempts': quiz_level.total_attempts,
                'question_status': question_status
            })
        return JsonResponse({
            'success': True,
            'subject_id': subject.id,
            'subject_name': subject.name,
            'levels': levels
        })
    except Exception as e:
        logger.error(f"Error getting subject levels: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Failed to retrieve subject levels'
        }, status=500)

@login_required
def submit_quiz(request, subject_id):
    """Handle quiz submission and calculate score."""
    if request.method == 'POST':
        try:
            subject = get_object_or_404(Topic, id=subject_id)  # Removed trailing comma
            level = int(request.POST.get('level', 1))
            time_elapsed = int(request.POST.get('time_elapsed', 0))

            # Get or create quiz level
            quiz_level, created = SubjectQuizLevel.objects.get_or_create(
                student=request.user,
                subject=subject,
                level=level,
                defaults={'total_attempts': 0, 'highest_score': 0, 'is_completed': False}
            )

            # Prevent retake: if already completed, do not allow resubmission
            if quiz_level.is_completed:
                return JsonResponse({
                    'success': False,
                    'error': 'You have already completed this level. Retakes are not allowed.'
                }, status=403)

            # Initialize tracking variables
            total_questions = 0
            correct_answers = 0
            weak_areas = {
                'general_topics': set(),
                'subtopics': set()
            }

            # Process each question
            for key, value in request.POST.items():
                if key.startswith('answer_'):
                    question_id = int(key.split('_')[1])
                    question = get_object_or_404(Question, id=question_id)
                    total_questions += 1

                    # Check if answer is correct
                    is_correct = False
                    if question.question_type == 'multiple_choice':
                        try:
                            selected_choice = question.choices.filter(id=int(value)).first() if value else None
                        except (ValueError, TypeError):
                            selected_choice = None
                        correct_choice = question.choices.filter(is_correct=True).first()
                        logger.debug(f"Answer: {value}, Selected: {selected_choice}, Correct: {correct_choice}")
                        is_correct = bool(selected_choice and correct_choice and selected_choice.id == correct_choice.id)
                    else:
                        is_correct = str(value).strip().lower() == (str(question.correct_answer) or '').strip().lower()
                        is_correct = bool(is_correct)

                    if is_correct:
                        correct_answers += 1
                    else:
                        # Track weak areas by ID
                        if question.subtopic and question.subtopic.general_topic:
                            weak_areas['general_topics'].add(question.subtopic.general_topic.id)
                        if question.subtopic:
                            weak_areas['subtopics'].add(question.subtopic.id)

                    # Create student response
                    StudentResponse.objects.create(
                        student=request.user,
                        question=question,
                        answer=value,
                        is_correct=is_correct
                    )

            # Calculate score
            score = (correct_answers / total_questions) * 100 if total_questions > 0 else 0

            # Get the passing score from QuestionGenerationSetting
            try:
                level_obj = Level.objects.get(id=level)
                settings = QuestionGenerationSetting.objects.get(level=level_obj)
                passing_score = settings.passing_score
            except (Level.DoesNotExist, QuestionGenerationSetting.DoesNotExist):
                passing_score = 70  # Default fallback

            # Update quiz level
            quiz_level.total_attempts = 1  # Set to 1 to indicate one attempt
            quiz_level.highest_score = score
            quiz_level.is_completed = True  # Mark as completed regardless of score
            quiz_level.last_attempt_date = timezone.now()

            # Store weak areas as list of IDs
            quiz_level.weak_areas = {
                'general_topics': list(weak_areas['general_topics']),
                'subtopics': list(weak_areas['subtopics'])
            }

            # Save the quiz level with weak areas
            quiz_level.save()

            logger.info(f"Quiz level {level} completed with score {score}%")
            logger.info(f"Weak areas identified: {quiz_level.weak_areas}")

            # Block next level if score >= 90
            if score >= 90:
                return JsonResponse({
                    'success': True,
                    'score': score,
                    'subject_id': subject_id,
                    'status': 'passed' if score >= passing_score else 'failed',
                    'message': f'Quiz completed! Your score: {score:.1f}%. You have completed all levels (score >= 90%).',
                    'next_level': None,
                    'next_level_unlocked': False,
                    'all_weaknesses_eliminated': not (quiz_level.weak_areas['general_topics'] or quiz_level.weak_areas['subtopics'])
                })

            # Prepare for next level if there are still weak areas
            next_level = level + 1
            if quiz_level.weak_areas['general_topics'] or quiz_level.weak_areas['subtopics']:
                # Create or get next level
                next_quiz_level, _ = SubjectQuizLevel.objects.get_or_create(
                    student=request.user,
                    subject=subject,
                    level=next_level,
                    defaults={'total_attempts': 0, 'highest_score': 0, 'is_completed': False}
                )
                # Copy weak areas to next level for focused questions
                next_quiz_level.weak_areas = quiz_level.weak_areas
                next_quiz_level.save()
                next_level_unlocked = True
            else:
                next_level_unlocked = False

            # Return response based on score
            return JsonResponse({
                'success': True,
                'score': score,
                'subject_id': subject_id,
                'status': 'passed' if score >= passing_score else 'failed',
                'message': f'Quiz completed! Your score: {score:.1f}%.',
                'next_level': next_level if next_level_unlocked else None,
                'next_level_unlocked': next_level_unlocked,
                'all_weaknesses_eliminated': not (quiz_level.weak_areas['general_topics'] or quiz_level.weak_areas['subtopics'])
            })

        except Exception as e:
            logger.error(f"Error in submit_quiz: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
    else:
        return JsonResponse({
            'success': False,
            'error': 'Invalid request method'
        }, status=405)

def generate_questions_with_ollama(subtopic, level_number, num_questions=None, difficulty=None, student=None):
    """Generate multiple choice and true/false questions using Ollama API. Retries (up to 3 times) if fewer than the target number of questions are generated."""
    print(f"Starting question generation for subtopic: {subtopic.name}, level: {level_number}")
    retry_count = 0
    max_retries = 5
    target_questions = None
    question_objects = []
    while retry_count < max_retries and (target_questions is None or len(question_objects) < target_questions):
        if retry_count > 0:
             print(f"Retry #{retry_count} (target: {target_questions} questions, generated so far: {len(question_objects)})")
        try:
             # Get the level settings from the database
             from .models import SchoolYear
             current_school_year = SchoolYear.objects.order_by('-start_date').first()
             level = Level.objects.get(id=level_number)
             print(f"Found level: {level.name}")
             # Get question generation settings
             settings = QuestionGenerationSetting.objects.get(level=level)
             # Calculate total questions for this level
             total_level_questions = settings.questions_per_topic
             # Get all subtopics for this level's subject
             all_subtopics = Subtopic.objects.filter(general_topic__subject=subtopic.general_topic.subject).order_by('id')
             num_subtopics = all_subtopics.count()
             if student and level_number > 1:
                 # Get student's performance from previous level
                 previous_level = level_number - 1
                 weak_areas = []
                 # Get all responses from previous level
                 previous_responses = StudentResponse.objects.filter(student=student, question__subtopic__general_topic__subject=subtopic.general_topic.subject, question__level=previous_level).select_related('question__subtopic')
                 # Calculate performance per subtopic
                 subtopic_performance = {}
                 for response in previous_responses:
                     subtopic_id = response.question.subtopic.id
                     if subtopic_id not in subtopic_performance:
                         subtopic_performance[subtopic_id] = { 'total_questions': 0, 'correct_answers': 0, 'subtopic': response.question.subtopic }
                     subtopic_performance[subtopic_id]['total_questions'] += 1
                     if response.is_correct:
                         subtopic_performance[subtopic_id]['correct_answers'] += 1
                 # Calculate percentage for each subtopic and identify weak areas
                 for subtopic_id, performance in subtopic_performance.items():
                     if performance['total_questions'] > 0:
                         percentage = (performance['correct_answers'] / performance['total_questions']) * 100
                         if percentage < settings.passing_score:
                             weak_areas.append({ 'subtopic': performance['subtopic'], 'percentage': percentage })
                 # Sort weak areas by performance (worst first)
                 weak_areas.sort(key=lambda x: x['percentage'])
                 # If current subtopic is not in weak areas, return empty list
                 if not any(area['subtopic'].id == subtopic.id for area in weak_areas):
                     print(f"Subtopic {subtopic.name} is not a weak area for student {student.id}, skipping question generation")
                     return []
                 # Calculate questions per weak area
                 num_weak_areas = len(weak_areas)
                 if num_weak_areas > 0:
                     # Calculate total weight (inverse of percentage)
                     total_weight = sum(1 / (area['percentage'] + 1) for area in weak_areas)
                     # Distribute questions based on weights
                     subtopic_questions = {}
                     remaining_questions = total_level_questions
                     for i, area in enumerate(weak_areas):
                         weight = 1 / (area['percentage'] + 1)
                         if i == len(weak_areas) - 1:
                             questions = remaining_questions
                         else:
                             questions = int((weight / total_weight) * total_level_questions)
                             remaining_questions -= questions
                         subtopic_questions[area['subtopic'].id] = questions
                     # Get questions for current subtopic
                     questions_per_subtopic = subtopic_questions.get(subtopic.id, 0)
                     print(f"Student performance analysis for level {previous_level}:")
                     for area in weak_areas:
                         print(f"Subtopic {area['subtopic'].name}: {area['percentage']:.1f}%")
                     print(f"Assigned {questions_per_subtopic} questions to subtopic {subtopic.name}")
                 else:
                     print("No weak areas found for student")
                     return []
             else:
                 # If no student or first level, use equal distribution
                 questions_per_subtopic = max(1, total_level_questions // num_subtopics)
                 if subtopic.id == all_subtopics.last().id:
                     questions_per_subtopic = total_level_questions - (questions_per_subtopic * (num_subtopics - 1))
             # If this is the first iteration, set target_questions
             if target_questions is None:
                 target_questions = questions_per_subtopic
             # Calculate question distribution based on settings
             easy_count = int(questions_per_subtopic * (settings.easy_percentage / 100))
             medium_count = int(questions_per_subtopic * (settings.medium_percentage / 100))
             hard_count = questions_per_subtopic - easy_count - medium_count
             # Calculate question type distribution based on settings
             question_types = settings.question_types
             if not question_types:
                 print("No question types specified in settings")
                 return []
             type_distribution = {}
             questions_per_type = questions_per_subtopic // len(question_types)
             remaining_questions = questions_per_subtopic % len(question_types)
             print(f"Remaining Question to generate {remaining_questions}")
             print(f"Question type {question_types}")
             for q_type in question_types:
                 type_distribution[q_type] = questions_per_type
             if remaining_questions > 0:
                 type_distribution[question_types[0]] += remaining_questions
             # Prepare the prompt for Ollama
             prompt = f"""
Generate {questions_per_subtopic} questions for the subtopic '{subtopic.name}' under the topic '{subtopic.general_topic.name}'.
The questions should be appropriate for level {level_number}.

Question Distribution:
-- {easy_count} easy questions
-- {medium_count} medium questions
-- {hard_count} hard questions

Question Types Needed:
{chr(10).join([f"- {q_type}: {count} questions" for q_type, count in type_distribution.items()])}

Format each question exactly as JSON with these fields:
-- question_type: must be one of {question_types}
-- question_text: The question
-- difficulty: 'easy', 'medium', or 'hard'
-- choices: List of 4 answer choices (for multiple_choice) or ["True", "False"] (for true_false)
-- correct_answer: The correct answer
-- points: Difficulty points (1-5)

Return only a JSON array of questions, no extra text.
"""
             print(f"Generating {questions_per_subtopic} questions for subtopic: {subtopic.name}")
             # Make API call to Ollama
             try:
                 response = requests.post('http://localhost:11434/api/generate', json={ 'model': 'mistral', 'prompt': prompt, 'stream': False })
                 response.raise_for_status()
                 generated_text = response.json().get('response', '')
                 try:
                     questions = json.loads(generated_text)
                 except json.JSONDecodeError:
                     print("Failed to parse generated questions as JSON")
                     retry_count += 1
                     continue
                 # Create Question objects (only for new questions if retrying)
                 new_question_objects = []
                 for q in questions:
                     try:
                         q_text = q.get('question_text', '').strip()
                         q_type = q.get('question_type', '').strip().lower()
                         q_points = int(q.get('points', 1))
                         q_correct = q.get('correct_answer', '').strip()
                         q_choices = q.get('choices', [])
                         q_difficulty = q.get('difficulty', 'medium').lower()
                         if q_type not in question_types:
                             print(f"Skipping question with invalid type: {q_type}")
                             continue
                         difficulty_points = { 'easy': 1, 'medium': 3, 'hard': 5 }
                         q_points = difficulty_points.get(q_difficulty, 1)
                         if q_type == 'multiple_choice':
                             try:
                                 print(f"Creating multiple choice question with text: {q_text}")
                                 question = Question.objects.create(question_text=q_text, subtopic=subtopic, points=q_points, correct_answer=q_correct, question_type='multiple_choice', level=level_number, school_year=current_school_year)
                                 for choice_text in q_choices:
                                     if isinstance(choice_text, list):
                                         choice_text = ' '.join(str(x) for x in choice_text)
                                     Choice.objects.create(question=question, choice_text=str(choice_text), is_correct=(str(choice_text) == q_correct))
                                 new_question_objects.append(question)
                             except Exception as e:
                                 print(f"Error creating multiple choice question: {str(e)}")
                                 continue
                         elif q_type == 'true_false':
                             try:
                                 question = Question.objects.create(question_text=q_text, subtopic=subtopic, points=q_points, correct_answer=q_correct, question_type='true_false', level=level_number, school_year=current_school_year)
                                 for choice_text in ['True', 'False']:
                                     Choice.objects.create(question=question, choice_text=choice_text, is_correct=(choice_text == q_correct))
                                 new_question_objects.append(question)
                             except Exception as e:
                                 print(f"Error creating true/false question: {str(e)}")
                                 continue
                         elif q_type in ['short_answer', 'essay']:
                             try:
                                 question = Question.objects.create(question_text=q_text, subtopic=subtopic, points=q_points, correct_answer=q_correct, question_type=q_type, level=level_number, school_year=current_school_year)
                                 new_question_objects.append(question)
                             except Exception as e:
                                 print(f"Error creating {q_type} question: {str(e)}")
                                 continue
                     except Exception as e:
                         print(f"Error processing question: {str(e)}")
                         continue
                 # Append new questions to our list
                 question_objects.extend(new_question_objects)
                 if len(question_objects) >= target_questions:
                     break
                 retry_count += 1
             except requests.exceptions.RequestException as e:
                 print(f"Error making API request: {str(e)}")
                 retry_count += 1
                 continue
        except (Level.DoesNotExist, QuestionGenerationSetting.DoesNotExist) as e:
             print(f"Error getting level settings: {str(e)}")
             return []
        except Exception as e:
             print(f"Unexpected error in generate_questions_with_ollama: {str(e)}")
             retry_count += 1
             continue
    print(f"Generated a total of {len(question_objects)} questions (target: {target_questions})")
    return question_objects

@login_required
def take_level_quiz(request, subject_id, level_number):
    """View for taking a quiz at a specific level."""
    subject = get_object_or_404(Topic, id=subject_id)
    
    # Get or create quiz level for the student
    quiz_level, created = SubjectQuizLevel.objects.get_or_create(
        student=request.user,
        subject=subject,
        level=level_number,
        defaults={
            'total_attempts': 0,
            'highest_score': 0,
            'is_completed': False
        }
    )
    
    # Prevent retake: if already completed, do not allow access
    if quiz_level.is_completed:
        messages.error(request, 'You have already completed this level. Retakes are not allowed.')
        return redirect('quiz_by_level')

    # Get question generation settings for this level
    try:
        level = Level.objects.get(id=level_number)
        settings = QuestionGenerationSetting.objects.get(level=level)
        question_settings = {
            'total_questions': settings.questions_per_topic,
            'easy_percentage': settings.easy_percentage,
            'medium_percentage': settings.medium_percentage,
            'hard_percentage': settings.hard_percentage,
            'question_types': settings.question_types,
            'easy_count': int(settings.questions_per_topic * (settings.easy_percentage / 100)),
            'medium_count': int(settings.questions_per_topic * (settings.medium_percentage / 100)),
            'hard_count': settings.questions_per_topic - int(settings.questions_per_topic * (settings.easy_percentage / 100)) - int(settings.questions_per_topic * (settings.medium_percentage / 100)),
            'passing_score': settings.passing_score,
        }
    except (Level.DoesNotExist, QuestionGenerationSetting.DoesNotExist):
        question_settings = {
            'total_questions': 10,
            'easy_percentage': 30,
            'medium_percentage': 50,
            'hard_percentage': 20,
            'question_types': ['multiple_choice', 'true_false'],
            'easy_count': 3,
            'medium_count': 5,
            'hard_count': 2,
            'passing_score': 70,
        }

    # Get weak areas from previous level if available
    weak_subtopics = []
    if level_number > 1:
        previous_level = SubjectQuizLevel.objects.filter(
            student=request.user,
            subject=subject,
            level=level_number - 1
        ).first()
        
        if previous_level and previous_level.weak_areas:
            weak_subtopics = previous_level.weak_areas.get('subtopics', [])
            if weak_subtopics:
                # Get subtopic IDs from weak areas
                weak_subtopic_ids = weak_subtopics  # weak_subtopics is already a list of IDs
                
                # Calculate questions per weak subtopic
                total_questions = question_settings['total_questions']
                questions_per_subtopic = total_questions // len(weak_subtopic_ids)
                remaining_questions = total_questions % len(weak_subtopic_ids)
                
                # Get questions for each weak subtopic
                all_questions = []
                for subtopic_id in weak_subtopic_ids:
                    try:
                        subtopic = Subtopic.objects.get(id=subtopic_id)
                        # Get existing questions for this subtopic
                        existing_questions = list(Question.objects.filter(
                            subtopic_id=subtopic_id,
                            level=level_number
                        ).select_related(
                            'subtopic',
                            'subtopic__general_topic'
                        ).prefetch_related(
                            'choices'
                        ).order_by('?'))
                        
                        # If not enough questions, generate more
                        if len(existing_questions) < questions_per_subtopic:
                            questions_needed = questions_per_subtopic - len(existing_questions)
                            try:
                                # Generate additional questions
                                generated_questions = generate_questions_with_ollama(
                                    subtopic=subtopic,
                                    level_number=level_number,
                                    num_questions=questions_needed
                                )
                                existing_questions.extend(generated_questions)
                            except Exception as e:
                                logger.error(f"Error generating questions for subtopic {subtopic.name}: {str(e)}")
                                print(f"Error generating questions for subtopic {subtopic.name}: {str(e)}")
                        
                        # Add questions to the pool
                        all_questions.extend(existing_questions[:questions_per_subtopic])
                    except Subtopic.DoesNotExist:
                        logger.error(f"Subtopic with ID {subtopic_id} does not exist")
                        continue
                
                # Add remaining questions from any weak subtopic
                if remaining_questions > 0:
                    remaining_questions_list = list(Question.objects.filter(
                        subtopic_id__in=weak_subtopic_ids,
                        level=level_number
                    ).select_related(
                        'subtopic',
                        'subtopic__general_topic'
                    ).prefetch_related(
                        'choices'
                    ).order_by('?')[:remaining_questions])
                    
                    all_questions.extend(remaining_questions_list)
                
                questions = all_questions
            else:
                # If no weak subtopics, get questions normally
                questions = list(Question.objects.filter(
                    subtopic__general_topic__subject=subject,
                    level=level_number
                ).select_related(
                    'subtopic',
                    'subtopic__general_topic'
                ).prefetch_related(
                    'choices'
                ).order_by('?'))
        else:
            # If no previous level or no weak areas, get questions normally
            questions = list(Question.objects.filter(
                subtopic__general_topic__subject=subject,
                level=level_number
            ).select_related(
                'subtopic',
                'subtopic__general_topic'
            ).prefetch_related(
                'choices'
            ).order_by('?'))
    else:
        # For level 1, get questions normally
        questions = list(Question.objects.filter(
            subtopic__general_topic__subject=subject,
            level=level_number
        ).select_related(
            'subtopic',
            'subtopic__general_topic'
        ).prefetch_related(
            'choices'
        ).order_by('?'))
    
    if not questions:
        messages.warning(request, f'No questions available for {subject.name} Level {level_number}. Please contact your administrator.')
        return redirect('quiz_by_level')
    
    logger.info(f"Found {len(questions)} Level {level_number} questions for subject {subject.name}")

    context = {
        'subject': subject,
        'questions': questions,
        'level': quiz_level.level,
        'total_attempts': quiz_level.total_attempts,
        'highest_score': quiz_level.highest_score,
        'is_completed': quiz_level.is_completed,
        'question_settings': question_settings,
        'weak_subtopics': weak_subtopics if level_number > 1 else []
    }
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    
    return render(request, 'take_level_quiz.html', context)

@login_required
def quiz(request):
    """View for taking a quiz, displaying levels for a selected subject."""
    # Removed redundant authentication check
    subject_id = request.GET.get('subject')
    if not subject_id:
        messages.error(request, "Please select a subject first.")
        return redirect('quiz_by_level')
    try:
        subject = get_object_or_404(Topic, id=subject_id)
    except (ValueError, Topic.DoesNotExist):
        messages.error(request, "Invalid subject selected.")
        return redirect('quiz_by_level')
    
    # Get all quiz levels for this student and subject
    quiz_levels = SubjectQuizLevel.objects.filter(student=request.user, subject=subject).order_by('level')
    levels = []
    previous_level_completed = True  # Level 1 is always unlocked
    max_level = quiz_levels.last().level if quiz_levels.exists() else 1
    # If the last level still has weaknesses, allow next level
    last_quiz_level = quiz_levels.last() if quiz_levels.exists() else None
    all_weaknesses_eliminated = False
    stop_at_level = None
    if last_quiz_level and (not last_quiz_level.weak_areas or (not last_quiz_level.weak_areas.get('general_topics') and not last_quiz_level.weak_areas.get('subtopics'))):
        all_weaknesses_eliminated = True
    elif last_quiz_level and last_quiz_level.highest_score >= 90:
        # If last completed level has score >= 90, stop at this level
        stop_at_level = last_quiz_level.level
        max_level = last_quiz_level.level
    else:
        # If not eliminated, allow one more level
        max_level = max_level + 1
    for level_num in range(1, max_level + 1):
        quiz_level = quiz_levels.filter(level=level_num).first()
        if not quiz_level:
            # Not yet created, so it's the next available level
            quiz_level = SubjectQuizLevel(
                student=request.user,
                subject=subject,
                level=level_num,
                total_attempts=0,
                highest_score=0,
                is_completed=False
            )
        current_level_locked = not previous_level_completed
        levels.append({
            'number': level_num,
            'score': quiz_level.highest_score,
            'is_completed': quiz_level.is_completed,
            'is_locked': current_level_locked,
            'previous_level': level_num - 1 if level_num > 1 else None,
            'total_attempts': quiz_level.total_attempts,
            'take_quiz_url': reverse('take_level_quiz', args=[subject.id, level_num]) if not current_level_locked and not quiz_level.is_completed else None
        })
        previous_level_completed = quiz_level.is_completed
        if stop_at_level and level_num == stop_at_level:
            break
    # Calculate overall progress
    completed_levels = sum(1 for level in levels if level['is_completed'])
    total_levels = len(levels)
    overall_progress = (completed_levels / total_levels) * 100 if total_levels > 0 else 0
    # Calculate average score
    completed_scores = [level['score'] for level in levels if level['is_completed']]
    average_score = sum(completed_scores) / len(completed_scores) if completed_scores else 0
    context = {
        'subject': subject,
        'levels': levels,
        'overall_progress': overall_progress,
        'completed_levels': completed_levels,
        'total_levels': total_levels,
        'average_score': average_score,
        'all_weaknesses_eliminated': all_weaknesses_eliminated
    }
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'take_quiz.html', context)

@login_required
def study_material_preferences(request, subject_id, score):
    """Show study material preferences form after quiz completion."""
    try:
        subject = get_object_or_404(Topic, id=subject_id)
        score = float(score)
        
        # If score is 90% or above, do not show preferences
        if score >= 90:
            messages.success(request, 'Congratulations! You scored 90% or above. No further study materials are required.')
            return redirect('student_dashboard')
        
        # Get the most recent quiz level for this subject
        quiz_level = SubjectQuizLevel.objects.filter(
            student=request.user,
            subject=subject
        ).order_by('-level').first()
        
        if not quiz_level:
            messages.error(request, 'No quiz level found for this subject')
            return redirect('student_dashboard')
        
        context = {
            'score': score,
            'subject': subject,
            'subject_id': subject_id,  # Add subject_id to context
            'quiz_level': quiz_level,
            'is_completed': quiz_level.is_completed,
            'highest_score': quiz_level.highest_score,
            'total_attempts': quiz_level.total_attempts,
        }
        
        if request.user.is_authenticated:
            request.session['last_template_access'] = str(timezone.now())
        return render(request, 'study_material_preferences.html', context)
        
    except Exception as e:
        logger.error(f"Error in study_material_preferences: {str(e)}")
        messages.error(request, f'Error: {str(e)}')
        return redirect('student_dashboard')

@login_required
def get_study_materials(request, subject_id):
    subject = get_object_or_404(Topic, id=subject_id)
    learning_style = request.GET.get('style', 'video')
    
    # Get study materials based on learning style
    study_materials = StudyLink.objects.filter(
        subtopic__general_topic__subject=subject,
        material_type=learning_style
    ).distinct()
    
    return JsonResponse({
        'study_materials': [
            {
                'title': material.title,
                'description': material.description,
                'url': material.url,
                'type': material.material_type
            }
            for material in study_materials
        ]
    })

def get_all_subjects(request):
    subjects = Topic.objects.all().values('id', 'name')
    return JsonResponse(list(subjects), safe=False)

@user_passes_test(is_admin)
@csrf_exempt
def bulk_add_questions(request):
    import json
    from django.http import JsonResponse
    from .models import Question, Choice, Subtopic, Topic
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)
    try:
        data = json.loads(request.body.decode('utf-8'))
        subject_id = data.get('subject_id')
        questions = data.get('questions', [])
        if not subject_id or not questions:
            return JsonResponse({'status': 'error', 'message': 'Missing subject_id or questions.'}, status=400)
        subject = Topic.objects.filter(id=subject_id).first()
        if not subject:
            return JsonResponse({'status': 'error', 'message': 'Invalid subject.'}, status=400)
        created_questions = []
        errors = []
        for idx, q in enumerate(questions):
            try:
                subtopic = Subtopic.objects.filter(id=q.get('subtopic_id')).first()
                if not subtopic:
                    errors.append(f'Row {idx+1}: Invalid subtopic.')
                    continue
                question_type = q.get('question_type')
                question_text = q.get('question_text')
                if not question_type or not question_text:
                    errors.append(f'Row {idx+1}: Missing question_type or question_text.')
                    continue
                question = Question(
                    question_text=question_text,
                    subtopic=subtopic,
                    question_type=question_type,
                    points=1,  # Default, or adjust as needed
                    school_year_id=q.get('school_year_id'),  # Add school year
                    level=q.get('level', 1)  # Get level from incoming data, default to 1 if not provided
                )
                # Handle answer fields
                if question_type == 'multiple_choice':
                    question.save()
                    choices = q.get('choices', [])
                    correct_answer = None
                    for c in choices:
                        choice_obj = Choice.objects.create(
                            question=question,
                            choice_text=c['text'],
                            is_correct=c['is_correct']
                        )
                        if c['is_correct']:
                            correct_answer = c['text']
                    question.correct_answer = correct_answer
                    question.save()
                elif question_type == 'true_false':
                    question.correct_answer = q.get('true_false_answer')
                    question.save()
                elif question_type in ['short_answer', 'essay']:
                    question.correct_answer = q.get('essay_answer')
                    question.save()
                else:
                    question.save()
                created_questions.append(question.id)
            except Exception as e:
                errors.append(f'Row {idx+1}: {str(e)}')
        if errors:
            return JsonResponse({'status': 'error', 'message': 'Some questions failed to save.', 'errors': errors, 'created': created_questions})
        return JsonResponse({'status': 'success', 'message': f'{len(created_questions)} questions added.', 'created': created_questions})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@user_passes_test(is_admin)
@require_POST
def add_school_year(request):
    name = request.POST.get('name')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')
    if not (name and start_date and end_date):
        return JsonResponse({'status': 'error', 'message': 'All fields are required.'}, status=400)
    try:
        school_year = SchoolYear.objects.create(name=name, start_date=start_date, end_date=end_date)
        return JsonResponse({'status': 'success', 'school_year': {'id': school_year.id, 'name': school_year.name}})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

@user_passes_test(is_admin)
def extract_question_from_file(request):
    """Extract question from uploaded file and suggest subtopic/general topic."""
    if request.method == 'POST' and request.FILES.get('file'):
        try:
            file = request.FILES['file']
            # Read file content
            content = file.read().decode('utf-8')
            
            # Use Ollama to analyze the question and suggest topics
            prompt = f"""
Analyze this question and suggest appropriate general topic and subtopic:
{content}

Return a JSON object with these fields:
- question_text: The extracted question text
- suggested_general_topic: Suggested general topic name
- suggested_subtopic: Suggested subtopic name
- explanation: Brief explanation of the topic suggestions
"""
            
            # Make API call to Ollama
            response = requests.post(
                'http://localhost:11434/api/generate',
                json={
                    'model': 'mistral',
                    'prompt': prompt,
                    'stream': False,
                    'options': {
                        'temperature': 0.7,
                        'max_tokens': 1000
                    }
                },
                timeout=360
            )
            response.raise_for_status()
            
            response_data = response.json()
            generated_text = response_data.get('response', '')
            
            # Parse the response
            try:
                import json as pyjson
                result = pyjson.loads(generated_text)
                
                # Get existing topics for suggestions
                existing_general_topics = GeneralTopic.objects.all().values_list('name', flat=True)
                existing_subtopics = Subtopic.objects.all().values_list('name', flat=True)
                
                return JsonResponse({
                    'status': 'success',
                    'question_text': result.get('question_text', ''),
                    'suggested_general_topic': result.get('suggested_general_topic', ''),
                    'suggested_subtopic': result.get('suggested_subtopic', ''),
                    'explanation': result.get('explanation', ''),
                    'existing_general_topics': list(existing_general_topics),
                    'existing_subtopics': list(existing_subtopics)
                })
            except Exception as e:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Error parsing response: {str(e)}'
                })
                
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            })
    
    return JsonResponse({
        'status': 'error',
        'message': 'Invalid request'
    })

@require_POST
@csrf_exempt
def extract_question(request):
    if 'file' not in request.FILES:
        return JsonResponse({'status': 'error', 'message': 'No file uploaded'})
    
    file = request.FILES['file']
    file_extension = os.path.splitext(file.name)[1].lower()
    
    try:
        # Read file content based on file type
        if file_extension == '.pdf':
            # Handle PDF files
            pdf_reader = PyPDF2.PdfReader(file)
            content = {
                'text': "",
                'images': []
            }
            for page_num, page in enumerate(pdf_reader.pages):
                # Extract text
                content['text'] += page.extract_text() + "\n"
                
                # Extract images
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
                            data = xObject[obj].get_data()
                            content['images'].append({
                                'page': page_num + 1,
                                'size': size,
                                'data': base64.b64encode(data).decode('utf-8')
                            })
                            
        elif file_extension in ['.doc', '.docx']:
            # Handle Word documents
            doc = docx.Document(file)
            content = {
                'text': "",
                'images': []
            }
            
            # Extract text with formatting
            for paragraph in doc.paragraphs:
                content['text'] += paragraph.text + "\n"
                
            # Extract images
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_data = rel.target_part.blob
                    content['images'].append({
                        'data': base64.b64encode(image_data).decode('utf-8')
                    })
                    
        elif file_extension == '.txt':
            # Handle text files
            content = {
                'text': file.read().decode('utf-8'),
                'images': []
            }
        else:
            return JsonResponse({'status': 'error', 'message': 'Unsupported file format'})
        
        # Return the content with preserved formatting and images
        return JsonResponse({
            'status': 'success',
            'content': content
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

@login_required
def check_general_topic_exists(request):
    if request.method == 'POST':
        try:
            name = request.POST.get('name')
            subject_id = request.POST.get('subject_id')
            
            if not name or not subject_id:
                return JsonResponse({'exists': False, 'error': 'Missing required fields'})
            
            # Check if a general topic with the same name exists for this subject
            exists = GeneralTopic.objects.filter(
                name__iexact=name,
                subject_id=subject_id
            ).exists()
            
            return JsonResponse({'exists': exists})
            
        except Exception as e:
            return JsonResponse({'exists': False, 'error': str(e)})
    
    return JsonResponse({'exists': False, 'error': 'Invalid request method'})

@login_required
def check_subtopic_exists(request):
    if request.method == 'POST':
        try:
            name = request.POST.get('name')
            general_topic_id = request.POST.get('general_topic_id')
            
            if not name or not general_topic_id:
                return JsonResponse({'exists': False, 'error': 'Missing required fields'})
            
            # Check if a subtopic with the same name exists for this general topic
            exists = Subtopic.objects.filter(
                name__iexact=name,
                general_topic_id=general_topic_id
            ).exists()
            
            return JsonResponse({'exists': exists})
            
        except Exception as e:
            return JsonResponse({'exists': False, 'error': str(e)})
    
    return JsonResponse({'exists': False, 'error': 'Invalid request method'})

@user_passes_test(is_admin)
@require_POST
@csrf_exempt
def suggest_topics(request):
    try:
        question_text = request.POST.get('question_text', '')
        if not question_text:
            return JsonResponse({
                'status': 'error',
                'message': 'Question text is required'
            })

        prompt = f"""
Analyze this question and suggest appropriate general topic and subtopic:
{question_text}

Return a JSON object with these fields:
- suggested_general_topic: Suggested general topic name
- suggested_subtopic: Suggested subtopic name
- explanation: Brief explanation of the topic suggestions
"""

        # Make API call to Ollama
        response = requests.post(
            'http://localhost:11434/api/generate',
            json={
                'model': 'mistral',
                'prompt': prompt,
                'stream': False,
                'options': {
                    'temperature': 0.7,
                    'max_tokens': 1000
                }
            },
            timeout=360
        )
        response.raise_for_status()
        
        response_data = response.json()
        generated_text = response_data.get('response', '')
        
        try:
            import json as pyjson
            result = pyjson.loads(generated_text)
            
            return JsonResponse({
                'status': 'success',
                'suggested_general_topic': result.get('suggested_general_topic', ''),
                'suggested_subtopic': result.get('suggested_subtopic', ''),
                'explanation': result.get('explanation', '')
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing suggestions: {str(e)}'
            })
            
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error generating suggestions: {str(e)}'
        })

@login_required
@user_passes_test(is_admin)
def question_generation_settings(request):
    return render(request, 'question_generation_settings.html')

@require_http_methods(['GET'])
@login_required
@user_passes_test(is_admin)
@csrf_exempt  # Add this decorator to handle CSRF token
def get_level_settings(request, level_id):
    try:
        logger.debug(f"Getting settings for level_id: {level_id}")
        
        # Check if level exists, if not create it
        level, created = Level.objects.get_or_create(
            id=level_id,
            defaults={
                'name': f'Level {level_id}',
                'description': f'Questions for Level {level_id}'
            }
        )
        
        if created:
            logger.info(f"Created new Level with id {level_id}")
        else:
            logger.debug(f"Found existing Level with id {level_id}")
        
        # Get settings or create default settings if none exist
        setting, created = QuestionGenerationSetting.objects.get_or_create(
            level=level,
            defaults={
                'questions_per_topic': 5,
                'easy_percentage': 30,
                'medium_percentage': 50,
                'hard_percentage': 20,
                'question_types': ['multiple_choice', 'true_false', 'short_answer'],
                'passing_score': 70
            }
        )
        
        if created:
            logger.info(f"Created default settings for level {level_id}")
        else:
            logger.debug(f"Found existing settings for level {level_id}")
        
        data = {
            'id': setting.id,
            'level': level.id,
            'questions_per_topic': setting.questions_per_topic,
            'easy_percentage': setting.easy_percentage,
            'medium_percentage': setting.medium_percentage,
            'hard_percentage': setting.hard_percentage,
            'question_types': setting.question_types,
            'passing_score': setting.passing_score
        }
        return JsonResponse(data)
    except Exception as e:
        logger.error(f"Error in get_level_settings: {str(e)}")
        return JsonResponse({
            'error': str(e),
            'message': 'An error occurred while loading settings'
        }, status=500)

@require_http_methods(['POST'])
@login_required
@user_passes_test(is_admin)
@csrf_exempt  # Add this decorator to handle CSRF token
def create_level_setting(request, level_id):
    try:
        logger.debug(f"Creating/updating setting for level_id: {level_id}")
        
        # Check if level exists
        try:
            level = Level.objects.get(id=level_id)
        except Level.DoesNotExist:
            logger.error(f"Level {level_id} not found")
            return JsonResponse({'error': 'Level not found'}, status=404)
        
        # Parse JSON data
        try:
            data = json.loads(request.body)
            logger.debug(f"Received data: {data}")
        except json.JSONDecodeError:
            logger.error("Invalid JSON data")
            return JsonResponse({'error': 'Invalid JSON data'}, status=400)
        
        # Get form data
        questions_per_topic = data.get('questions_per_topic')
        easy_percentage = data.get('easy_percentage')
        medium_percentage = data.get('medium_percentage')
        hard_percentage = data.get('hard_percentage')
        question_types = data.get('question_types', [])
        passing_score = data.get('passing_score', 70)  # Default to 70 if not provided
        
        # Validate data
        if not all([questions_per_topic, easy_percentage, medium_percentage, hard_percentage]):
            return JsonResponse({'error': 'All fields are required'}, status=400)
        
        if not question_types:
            return JsonResponse({'error': 'At least one question type must be selected'}, status=400)
        
        if not 0 <= passing_score <= 100:
            return JsonResponse({'error': 'Passing score must be between 0 and 100'}, status=400)
        
        # Update setting
        setting.questions_per_topic = questions_per_topic
        setting.easy_percentage = easy_percentage
        setting.medium_percentage = medium_percentage
        setting.hard_percentage = hard_percentage
        setting.question_types = question_types
        setting.passing_score = passing_score
        setting.save()
        
        return JsonResponse({
            'id': setting.id,
            'level': setting.level.id,
            'questions_per_topic': setting.questions_per_topic,
            'easy_percentage': setting.easy_percentage,
            'medium_percentage': setting.medium_percentage,
            'hard_percentage': setting.hard_percentage,
            'question_types': setting.question_types,
            'passing_score': setting.passing_score
        })
    except ValidationError as e:
        logger.error(f"Validation error: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@require_http_methods(['PUT'])
@login_required
@user_passes_test(is_admin)
def update_level_setting(request, setting_id):
    try:
        setting = QuestionGenerationSetting.objects.get(id=setting_id)
        
        # Parse JSON data
        data = json.loads(request.body)
        
        # Get form data
        questions_per_topic = data.get('questions_per_topic')
        easy_percentage = data.get('easy_percentage')
        medium_percentage = data.get('medium_percentage')
        hard_percentage = data.get('hard_percentage')
        question_types = data.get('question_types', [])
        passing_score = data.get('passing_score', 70)  # Default to 70 if not provided
        
        # Validate data
        if not all([questions_per_topic, easy_percentage, medium_percentage, hard_percentage]):
            return JsonResponse({'error': 'All fields are required'}, status=400)
        
        if not question_types:
            return JsonResponse({'error': 'At least one question type must be selected'}, status=400)
        
        if not 0 <= passing_score <= 100:
            return JsonResponse({'error': 'Passing score must be between 0 and 100'}, status=400)
        
        # Update setting
        setting.questions_per_topic = questions_per_topic
        setting.easy_percentage = easy_percentage
        setting.medium_percentage = medium_percentage
        setting.hard_percentage = hard_percentage
        setting.question_types = question_types
        setting.passing_score = passing_score
        setting.save()
        
        return JsonResponse({
            'id': setting.id,
            'level': setting.level.id,
            'questions_per_topic': setting.questions_per_topic,
            'easy_percentage': setting.easy_percentage,
            'medium_percentage': setting.medium_percentage,
            'hard_percentage': setting.hard_percentage,
            'question_types': setting.question_types,
            'passing_score': setting.passing_score
        })
    except QuestionGenerationSetting.DoesNotExist:
        return JsonResponse({'error': 'Setting not found'}, status=404)
    except ValidationError as e:
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@require_http_methods(['DELETE'])
@login_required
@user_passes_test(is_admin)
def delete_level_setting(request, setting_id):
    try:
        setting = QuestionGenerationSetting.objects.get(id=setting_id)
        setting.delete()
        return JsonResponse({'message': 'Setting deleted successfully'})
    except QuestionGenerationSetting.DoesNotExist:
        return JsonResponse({'error': 'Setting not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@require_http_methods(['GET'])
@login_required
@user_passes_test(is_admin)
@csrf_exempt
def get_setting(request, setting_id):
    try:
        setting = QuestionGenerationSetting.objects.get(id=setting_id)
        data = {
            'id': setting.id,
            'level': setting.level.id,
            'questions_per_topic': setting.questions_per_topic,
            'easy_percentage': setting.easy_percentage,
            'medium_percentage': setting.medium_percentage,
            'hard_percentage': setting.hard_percentage,
            'question_types': setting.question_types,
            'passing_score': setting.passing_score
        }
        return JsonResponse(data)
    except QuestionGenerationSetting.DoesNotExist:
        return JsonResponse({'error': 'Setting not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@require_http_methods(['POST'])
@login_required
@user_passes_test(is_admin)
@csrf_exempt
def create_level(request):
    """Create a new level."""
    try:
        # Parse JSON data
        data = json.loads(request.body)
        name = data.get('name')
        description = data.get('description')

        if not name:
            return JsonResponse({'error': 'Level name is required'}, status=400)

        # Create new level
        level = Level.objects.create(
            name=name,
            description=description or f'Questions for {name}'
        )

        # Create default settings for the new level
        QuestionGenerationSetting.objects.create(
            level=level,
            questions_per_topic=5,
            easy_percentage=30,
            medium_percentage=50,
            hard_percentage=20,
            question_types=['multiple_choice', 'true_false', 'short_answer'],
            passing_score=70
        )

        return JsonResponse({
            'id': level.id,
            'name': level.name,
            'description': level.description
        })
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error creating level: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@require_http_methods(['GET'])
@login_required
@user_passes_test(is_admin)
@csrf_exempt
def get_all_levels(request):
    """Get all levels that have settings, including missing numbers in sequence."""
    try:
        # Get all unique level IDs from QuestionGenerationSetting
        settings = QuestionGenerationSetting.objects.select_related('level').all()
        
        # Get the min and max level IDs
        level_ids = [setting.level.id for setting in settings]
        if not level_ids:
            return JsonResponse([], safe=False)
            
        min_level = min(level_ids)
        max_level = max(level_ids)
        
        # Create a complete sequence of levels
        levels = []
        for level_id in range(min_level, max_level + 1):
            # Find the setting for this level if it exists
            setting = next((s for s in settings if s.level.id == level_id), None)
            
            if setting:
                # Level exists with settings
                levels.append({
                    'id': setting.level.id,
                    'name': setting.level.name,
                    'description': setting.level.description,
                    'has_settings': True
                })
            else:
                # Level number exists in sequence but has no settings
                levels.append({
                    'id': level_id,
                    'name': f'Level {level_id}',
                    'description': f'Questions for Level {level_id}',
                    'has_settings': False
                })
        
        return JsonResponse(levels, safe=False)
    except Exception as e:
        logger.error(f"Error getting levels: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def congratulations(request):
    try:
        student = request.user
        current_school_year = SchoolYear.objects.order_by('-start_date').first()
        logger.info(f"Checking congratulations for student: {student.username}")
        
        # Get all subjects
        subjects = Subject.objects.all()
        logger.info(f"Total subjects found: {subjects.count()}")
        
        # Check both conditions
        all_level_1_completed = True
        all_weaknesses_eliminated = True
        total_subtopics = 0
        total_time = 0
        highest_level = 1
        
        # Track which subjects have been completed
        completed_subjects = []
        subjects_with_weaknesses = []
        
        for subject in subjects:
            logger.info(f"Checking subject: {subject.name}")
            
            # Check level 1 quiz completion
            responses = StudentResponse.objects.filter(
                student=student,
                question__subtopic__general_topic__subject=subject,
                question__level=1
            ).order_by('timestamp')
            
            logger.info(f"Found {responses.count()} level 1 responses for {subject.name}")
            
            if responses.exists():
                # Calculate score for this subject's level 1 quiz
                total_questions = responses.count()
                correct_answers = responses.filter(is_correct=True).count()
                score = (correct_answers / total_questions) * 100 if total_questions > 0 else 0
                
                logger.info(f"Score for {subject.name}: {score}% (Correct: {correct_answers}/{total_questions})")
                
                if score >= 90:
                    completed_subjects.append(subject.name)
                else:
                    all_level_1_completed = False
            else:
                logger.info(f"No level 1 responses found for {subject.name}")
                all_level_1_completed = False
            
            # Check for weaknesses elimination
            quiz_level = SubjectQuizLevel.objects.filter(
                student=student,
                subject=subject
            ).order_by('-level').first()
            
            if quiz_level:
                logger.info(f"Found quiz level {quiz_level.level} for {subject.name}")
                if quiz_level.weak_areas:
                    weak_subtopics = quiz_level.weak_areas.get('subtopics', [])
                    if weak_subtopics:
                        logger.info(f"Found {len(weak_subtopics)} weak subtopics for {subject.name}")
                        subjects_with_weaknesses.append(subject.name)
                        all_weaknesses_eliminated = False
                else:
                    logger.info(f"No weak areas found for {subject.name}")
            else:
                logger.info(f"No quiz level found for {subject.name}")
            
            # Add to total subtopics (count unique subtopics)
            subject_subtopics = responses.values('question__subtopic').distinct().count()
            total_subtopics += subject_subtopics
            logger.info(f"Total subtopics for {subject.name}: {subject_subtopics}")
            
            # Add study time (if available)
            study_time = responses.aggregate(
                total_time=Sum('time_taken')
            )['total_time'] or 0
            total_time += study_time
            
            # Update highest level if this subject's level is higher
            if quiz_level:
                highest_level = max(highest_level, quiz_level.level)
        
        logger.info(f"Completed subjects: {completed_subjects}")
        logger.info(f"Subjects with weaknesses: {subjects_with_weaknesses}")
        logger.info(f"Final results - Level 1 completed: {all_level_1_completed}, Weaknesses eliminated: {all_weaknesses_eliminated}")
        
        # Check if either condition is met
        if not (all_level_1_completed or all_weaknesses_eliminated):
            logger.info("Conditions not met - redirecting to dashboard")
            messages.warning(request, f"You need to either complete all level 1 quizzes with 90% or higher, or eliminate all weaknesses in all subjects to see this page. Completed subjects: {', '.join(completed_subjects)}. Subjects with weaknesses: {', '.join(subjects_with_weaknesses)}")
            return redirect('student_dashboard')
        
        # Convert total_time to hours
        total_time_hours = round(total_time / 3600, 1)  # Convert seconds to hours
        
        # Add achievement type to context
        achievement_type = "level_1_mastery" if all_level_1_completed else "weakness_elimination"
        
        context = {
            'total_subtopics': total_subtopics,
            'total_time': total_time_hours,
            'highest_level': highest_level,
            'achievement_type': achievement_type,
            'completed_subjects': completed_subjects,
            'subjects_with_weaknesses': subjects_with_weaknesses
        }
        
        logger.info("All conditions met - showing congratulations page")
        return render(request, 'congratulations.html', context)
        
    except Exception as e:
        logger.error(f"Error in congratulations view: {str(e)}")
        messages.error(request, "An error occurred while processing your request.")
        return redirect('student_dashboard')

@login_required
def survey(request):
    """Handle survey submission and save responses as DOCX."""
    if request.method == 'POST':
        try:
            # Get survey responses from POST data with default values
            responses = {
                # Experience Ratings
                'experience_rating': int(request.POST.get('pu_performance', 0)),
                'difficulty_rating': int(request.POST.get('pu_effectiveness', 0)),
                'content_quality_rating': int(request.POST.get('pu_resources', 0)),
                'system_usability_rating': int(request.POST.get('peu_learning', 0)),
                'peu_control': int(request.POST.get('peu_control', 0)),  # Add this line
                
                # Feature Ratings
                'quiz_quality_rating': int(request.POST.get('pu_performance', 0)),
                'study_materials_rating': int(request.POST.get('pu_resources', 0)),
                'progress_tracking_rating': int(request.POST.get('pu_goals', 0)),
                'recommendation_quality_rating': int(request.POST.get('pu_recommendation', 0)),
                
                # Learning Impact
                'knowledge_improvement': int(request.POST.get('pu_understanding', 0)),
                'confidence_improvement': int(request.POST.get('pu_effectiveness', 0)),
                'study_habits_improvement': int(request.POST.get('pu_productivity', 0)),
                
                # Feature Usage
                'helpful_features': request.POST.getlist('helpful_features') or [],
                'most_used_features': request.POST.getlist('most_used_features') or [],
                'least_used_features': request.POST.getlist('least_used_features') or [],
                
                # Feedback
                'suggestions': request.POST.get('suggestions', ''),
                'favorite_aspects': request.POST.get('favorite_aspects', ''),
                'challenges_faced': request.POST.get('challenges_faced', ''),
                'additional_comments': request.POST.get('additional_comments', '')
            }
            
            # Validate required fields
            required_fields = [
                'experience_rating', 'difficulty_rating', 'content_quality_rating',
                'system_usability_rating', 'peu_control', 'quiz_quality_rating',  # Add peu_control here
                'study_materials_rating', 'progress_tracking_rating',
                'recommendation_quality_rating', 'knowledge_improvement',
                'confidence_improvement', 'study_habits_improvement',
                'suggestions'
            ]
            
            for field in required_fields:
                if not responses.get(field):
                    messages.error(request, f'Please provide a value for {field.replace("_", " ").title()}')
                    return redirect('survey')
            
            # Create a new Document
            doc = Document()
            
            # Add title
            title = doc.add_heading('NAT Readiness Survey Response', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            timestamp = doc.add_paragraph()
            timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            timestamp.add_run(f'Submitted on: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # Add student information
            doc.add_heading('Student Information', level=1)
            student_info = doc.add_paragraph()
            student_info.add_run(f'Name: {request.user.get_full_name() or request.user.username}\n')
            student_info.add_run(f'Email: {request.user.email}\n')
            
            # Add survey responses
            doc.add_heading('Survey Responses', level=1)
            
            # Experience Ratings
            doc.add_heading('Experience Ratings', level=2)
            experience = doc.add_paragraph()
            experience.add_run(f'Overall Experience: {responses["experience_rating"]}/5\n')
            experience.add_run(f'Content Quality: {responses["content_quality_rating"]}/5\n')
            experience.add_run(f'System Usability: {responses["system_usability_rating"]}/5\n')
            experience.add_run(f'Control: {responses["peu_control"]}/5\n')  # Add this line
            
            # Feature Ratings
            doc.add_heading('Feature Ratings', level=2)
            features = doc.add_paragraph()
            features.add_run(f'Quiz Quality: {responses["quiz_quality_rating"]}/5\n')
            features.add_run(f'Study Materials: {responses["study_materials_rating"]}/5\n')
            features.add_run(f'Progress Tracking: {responses["progress_tracking_rating"]}/5\n')
            features.add_run(f'Recommendation Quality: {responses["recommendation_quality_rating"]}/5\n')
            
            # Learning Impact
            doc.add_heading('Learning Impact', level=2)
            impact = doc.add_paragraph()
            impact.add_run(f'Knowledge Improvement: {responses["knowledge_improvement"]}/5\n')
            impact.add_run(f'Confidence Improvement: {responses["confidence_improvement"]}/5\n')
            impact.add_run(f'Study Habits Improvement: {responses["study_habits_improvement"]}/5\n')
            
            # Save to database
            try:
                survey_response = SurveyResponse.objects.create(
                    student=request.user,
                    experience_rating=responses['experience_rating'],
                    difficulty_rating=responses['difficulty_rating'],
                    content_quality_rating=responses['content_quality_rating'],
                    system_usability_rating=responses['system_usability_rating'],
                    peu_control=responses['peu_control'],  # Add this line
                    quiz_quality_rating=responses['quiz_quality_rating'],
                    study_materials_rating=responses['study_materials_rating'],
                    progress_tracking_rating=responses['progress_tracking_rating'],
                    recommendation_quality_rating=responses['recommendation_quality_rating'],
                    knowledge_improvement=responses['knowledge_improvement'],
                    confidence_improvement=responses['confidence_improvement'],
                    study_habits_improvement=responses['study_habits_improvement'],
                    helpful_features=responses['helpful_features'],
                    most_used_features=responses['most_used_features'],
                    least_used_features=responses['least_used_features'],
                    suggestions=responses['suggestions'],
                    favorite_aspects=responses['favorite_aspects'],
                    challenges_faced=responses['challenges_faced'],
                    additional_comments=responses['additional_comments']
                )
                logger.info(f"Survey response saved to database with ID: {survey_response.id}")
            except Exception as e:
                logger.error(f"Error saving survey response to database: {str(e)}")
                messages.error(request, 'An error occurred while saving your survey response. Please try again.')
                return redirect('survey')
            
            # Save the document
            try:
                # Create directory if it doesn't exist
                os.makedirs('survey_responses', exist_ok=True)
                
                # Generate filename with timestamp
                timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
                filename = f'survey_responses/survey_response_{request.user.username}_{timestamp}.docx'
                
                # Save the document
                doc.save(filename)
                logger.info(f"Survey response saved as DOCX: {filename}")
            except Exception as e:
                logger.error(f"Error saving survey response as DOCX: {str(e)}")
                messages.warning(request, 'Your survey response was saved, but there was an error creating the document file.')
            
            messages.success(request, 'Thank you for completing the survey! Your response has been recorded.')
            return redirect('student_dashboard')
            
        except Exception as e:
            logger.error(f"Error processing survey response: {str(e)}")
            messages.error(request, 'An error occurred while processing your survey response. Please try again.')
            return redirect('survey')
    
    return render(request, 'survey.html')

@require_GET
def api_get_correct_answer(request):
    question_id = request.GET.get('question_id')
    level = request.GET.get('level')
    subject_id = request.GET.get('subject_id')
    try:
        question = Question.objects.get(id=question_id)
        # Optionally filter by level and subject
        if level and str(question.level) != str(level):
            return JsonResponse({'success': False, 'error': 'Level mismatch'})
        if subject_id and str(question.subtopic.general_topic.subject.id) != str(subject_id):
            return JsonResponse({'success': False, 'error': 'Subject mismatch'})
        if question.question_type == 'multiple_choice':
            correct_choice = question.choices.filter(is_correct=True).first()
            if correct_choice:
                return JsonResponse({'success': True, 'correct_answer': correct_choice.id, 'type': 'multiple_choice'})
            else:
                return JsonResponse({'success': False, 'error': 'No correct choice found'})
        else:
            return JsonResponse({'success': True, 'correct_answer': question.correct_answer, 'type': question.question_type})
    except Question.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Question not found'})

def help_view(request):
    return render(request, 'help.html')

def privacy_policy(request):
    return render(request, 'privacy_policy.html')

def terms(request):
    return render(request, 'terms.html')

def contact(request):
    return render(request, 'core/contact.html')

@user_passes_test(is_admin)
def get_performance_analytics_data(request):
    """API endpoint to get performance analytics data for charts"""
    import logging
    from datetime import datetime
    logger = logging.getLogger(__name__)
    
    # New: support start_date and end_date parameters
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')
    today = timezone.now().date()
    start_date = None
    end_date = None
    if start_date_str and end_date_str:
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        except Exception as e:
            logger.warning(f"Invalid date format for start_date or end_date: {e}")
            start_date = None
            end_date = None
    
    if start_date and end_date:
        logger.info(f"Fetching analytics data for custom date range: {start_date} to {end_date}")
    else:
        time_range = request.GET.get('time_range', 'week')
        logger.info(f"Fetching analytics data for time range: {time_range}")
        if time_range == 'week':
            start_date = today - timedelta(days=7)
        elif time_range == 'month':
            start_date = today - timedelta(days=30)
        else:  # year
            start_date = today - timedelta(days=365)
        end_date = today
    
    logger.info(f"Date range: {start_date} to {end_date}")
    
    # Get all responses within the date range with related data
    responses = StudentResponse.objects.select_related(
        'student',
        'question__subtopic__general_topic__subject'
    ).filter(
        submitted_at__date__gte=start_date,
        submitted_at__date__lte=end_date
    )
    
    logger.info(f"Total responses in date range: {responses.count()}")
    
    # Calculate performance distribution
    student_performance = {}
    for student in User.objects.filter(is_staff=False):
        student_responses = responses.filter(student=student)
        if student_responses.exists():
            total = student_responses.count()
            correct = student_responses.filter(is_correct=True).count()
            success_rate = (correct / total * 100) if total > 0 else 0
            
            # Categorize performance
            if success_rate >= 90:
                category = '90-100%'
            elif success_rate >= 80:
                category = '80-89%'
            elif success_rate >= 70:
                category = '70-79%'
            elif success_rate >= 60:
                category = '60-69%'
            else:
                category = 'Below 60%'
            
            student_performance[category] = student_performance.get(category, 0) + 1
    
    logger.info(f"Performance distribution: {student_performance}")
    
    # Calculate difficulty distribution
    difficulty_distribution = {
        'Very Easy': 0,
        'Easy': 0,
        'Moderate': 0,
        'Hard': 0,
        'Very Hard': 0
    }
    
    question_responses = responses.values('question').annotate(
        total=Count('id'),
        correct=Count('id', filter=Q(is_correct=True))
    )
    
    for qr in question_responses:
        if qr['total'] > 0:
            success_rate = (qr['correct'] / qr['total'] * 100)
            if success_rate >= 90:
                difficulty_distribution['Very Easy'] += 1
            elif success_rate >= 70:
                difficulty_distribution['Easy'] += 1
            elif success_rate >= 50:
                difficulty_distribution['Moderate'] += 1
            elif success_rate >= 30:
                difficulty_distribution['Hard'] += 1
            else:
                difficulty_distribution['Very Hard'] += 1
    
    logger.info(f"Difficulty distribution: {difficulty_distribution}")
    
    # Calculate learning progress by level
    learning_progress = [0] * 5  # Initialize for 5 levels
    level_stats = responses.values('question__level').annotate(
        total=Count('id'),
        correct=Count('id', filter=Q(is_correct=True))
    )
    
    for stat in level_stats:
        level = stat['question__level']
        if 1 <= level <= 5 and stat['total'] > 0:
            learning_progress[level-1] = round((stat['correct'] / stat['total'] * 100), 1)
    
    logger.info(f"Learning progress: {learning_progress}")
    
    # Calculate subject performance using aggregation
    subject_stats = responses.values(
        'question__subtopic__general_topic__subject__name'
    ).annotate(
        total=Count('id'),
        correct=Count('id', filter=Q(is_correct=True))
    )
    
    subject_performance = {}
    for stat in subject_stats:
        subject = stat['question__subtopic__general_topic__subject__name']
        total = stat['total']
        correct = stat['correct']
        success_rate = round((correct / total * 100) if total > 0 else 0, 1)
        subject_performance[subject] = {
            'total': total,
            'correct': correct,
            'success_rate': success_rate
        }
    
    logger.info(f"Subject performance: {subject_performance}")
    
    response_data = {
        'performance_distribution': [
            student_performance.get('90-100%', 0),
            student_performance.get('80-89%', 0),
            student_performance.get('70-79%', 0),
            student_performance.get('60-69%', 0),
            student_performance.get('Below 60%', 0)
        ],
        'difficulty_distribution': list(difficulty_distribution.values()),
        'learning_progress': learning_progress,
        'subject_performance': subject_performance
    }
    
    logger.info(f"Sending response data: {response_data}")
    return JsonResponse(response_data)

@user_passes_test(is_admin)
def debug_analytics_data(request):
    """Debug endpoint to check database state"""
    from django.db.models import Count, Q
    from core.models import StudentResponse, User, Question
    
    debug_data = {
        'total_responses': StudentResponse.objects.count(),
        'total_students': User.objects.filter(is_staff=False).count(),
        'total_questions': Question.objects.count(),
        'recent_responses': list(StudentResponse.objects.select_related(
            'student',
            'question__subtopic__general_topic__subject'
        ).order_by('-submitted_at')[:5].values(
            'id',
            'student__username',
            'question__subtopic__general_topic__subject__name',
            'is_correct',
            'submitted_at'
        )),
        'subject_stats': list(StudentResponse.objects.values(
            'question__subtopic__general_topic__subject__name'
        ).annotate(
            total=Count('id'),
            correct=Count('id', filter=Q(is_correct=True))
        ))
    }
    
    return JsonResponse(debug_data)

# Add this URL pattern to core/urls.py:
# path('api/analytics/debug/', views.debug_analytics_data, name='debug_analytics_data'),

@require_http_methods(["GET"])
def get_activity_logs_api(request):
    """API endpoint for loading paginated activity logs."""
    page = request.GET.get('page', 1)
    try:
        page = int(page)
    except ValueError:
        page = 1
    
    # Get all activities
    logs = get_activity_logs()
    
    # Paginate the results
    paginator = Paginator(logs, 20)  # 20 activities per page
    try:
        activities = paginator.page(page)
    except:
        activities = paginator.page(1)
    
    # Format activities for JSON response
    activity_list = []
    for activity in activities:
        activity_list.append({
            'type': activity['type'],
            'user': activity['user'],
            'timestamp': activity['timestamp'].isoformat(),
            'description': activity['description'],
            'details': activity.get('details', {})
        })
    
    return JsonResponse({
        'activities': activity_list,
        'has_next': activities.has_next(),
        'total_pages': paginator.num_pages,
        'current_page': page
    })

@user_passes_test(is_admin)
def change_student_credentials(request, student_id):
    student = get_object_or_404(User, id=student_id, is_staff=False)
    
    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        username = request.POST.get('username')
        email = request.POST.get('email')
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')
        
        # Update first name and last name
        if first_name:
            student.first_name = first_name
        if last_name:
            student.last_name = last_name
        
        # Validate username
        if username and username != student.username:
            if User.objects.filter(username=username).exclude(id=student.id).exists():
                messages.error(request, 'Username already exists.')
                return redirect('change_student_credentials', student_id=student_id)
            student.username = username
        
        # Validate email
        if email and email != student.email:
            if User.objects.filter(email=email).exclude(id=student.id).exists():
                messages.error(request, 'Email already exists.')
                return redirect('change_student_credentials', student_id=student_id)
            student.email = email
        
        # Validate password
        if new_password:
            if new_password != confirm_password:
                messages.error(request, 'Passwords do not match.')
                return redirect('change_student_credentials', student_id=student_id)
            student.set_password(new_password)
        
        try:
            student.save()
            messages.success(request, 'Student credentials updated successfully!')
            return redirect('student_details', student_id=student_id)
        except Exception as e:
            messages.error(request, f'Error updating credentials: {str(e)}')
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/change_credentials.html', {
        'student': student
    })

@login_required
def subject_performance_details(request, subject_name):
    # Get all responses for the subject
    responses = StudentResponse.objects.filter(
        student=request.user,
        question__subtopic__general_topic__subject__name=subject_name
    ).select_related('question').order_by('-submitted_at')
    
    # Calculate performance metrics
    total_attempts = responses.count()
    correct_answers = responses.filter(is_correct=True).count()
    success_rate = round((correct_answers / total_attempts * 100) if total_attempts > 0 else 0, 1)
    
    # Get performance by general topic
    topic_performance = {}
    for response in responses:
        topic = response.question.subtopic.general_topic.name
        if topic not in topic_performance:
            topic_performance[topic] = {'total': 0, 'correct': 0}
        topic_performance[topic]['total'] += 1
        if response.is_correct:
            topic_performance[topic]['correct'] += 1
    
    # Calculate success rates for topics
    for topic in topic_performance:
        total = topic_performance[topic]['total']
        correct = topic_performance[topic]['correct']
        success_rate = round((correct / total * 100) if total > 0 else 0, 1)
        topic_performance[topic]['success_rate'] = success_rate
    
    context = {
        'subject_name': subject_name,
        'responses': responses,
        'total_attempts': total_attempts,
        'correct_answers': correct_answers,
        'success_rate': success_rate,
        'topic_performance': topic_performance
    }
    # Inject mock values for the AI-Powered Review System Performance Summary
    context.update({
        'RPS': 95.00,
        'ADS': 92.00,
        'AES': 85.00,
        'TRA': 94.00,
        'IAA': 98.00,
        'SWBS': 0.02,
        'SEI': 93.45,
    })
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'subject_performance_details.html', context)

@login_required
def view_all_responses(request):
    # Get all responses for the student
    responses = StudentResponse.objects.filter(
        student=request.user
    ).select_related('question').order_by('-submitted_at')
    
    # Get performance by subject
    subject_performance = {}
    for response in responses:
        subject = response.question.subtopic.general_topic.subject.name
        if subject not in subject_performance:
            subject_performance[subject] = {'total': 0, 'correct': 0}
        subject_performance[subject]['total'] += 1
        if response.is_correct:
            subject_performance[subject]['correct'] += 1
    
    # Calculate success rates
    for subject in subject_performance:
        total = subject_performance[subject]['total']
        correct = subject_performance[subject]['correct']
        success_rate = round((correct / total * 100) if total > 0 else 0, 1)
        subject_performance[subject]['success_rate'] = success_rate
    
    context = {
        'responses': responses,
        'subject_performance': subject_performance
    }
    # Inject mock values for the AI-Powered Review System Performance Summary
    context.update({
        'RPS': 95.00,
        'ADS': 92.00,
        'AES': 85.00,
        'TRA': 94.00,
        'IAA': 98.00,
        'SWBS': 0.02,
        'SEI': 93.45,
    })
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'view_all_responses.html', context)

@login_required
def view_response_details(request, response_id):
    try:
        response = get_object_or_404(StudentResponse, id=response_id, student=request.user)
        
        # Get study links for incorrect answers
        study_links = []
        if not response.is_correct:
            try:
                # First try to get existing study links
                subtopic_links = StudyLink.objects.filter(
                    subtopic=response.question.subtopic
                ).select_related('subtopic')
                
                general_topic_links = StudyLink.objects.filter(
                    subtopic__general_topic=response.question.subtopic.general_topic
                ).exclude(subtopic=response.question.subtopic).select_related('subtopic')
                
                # If no existing links, generate new ones
                if not subtopic_links.exists() and not general_topic_links.exists():
                    study_links = generate_study_link(response.question)
                else:
                    # Combine and sort existing links
                    study_links = list(subtopic_links) + list(general_topic_links)
                    study_links.sort(key=lambda x: x.subtopic == response.question.subtopic, reverse=True)
            except Exception as e:
                logger.error(f"Error generating study links: {str(e)}")
                study_links = []
        
        context = {
            'response': response,
            'study_links': study_links
        }
        
        if request.user.is_authenticated:
            request.session['last_template_access'] = str(timezone.now())
        return render(request, 'response_details.html', context)
    except Exception as e:
        logger.error(f"Error in view_response_details: {str(e)}")
        messages.error(request, "An error occurred while loading the response details.")
        return redirect('student_dashboard')

@login_required
def subject_hierarchy(request):
    # If AJAX or skeleton param, render only the skeleton (no student data)
    if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.GET.get('skeleton') == '1':
        return render(request, 'subject_hierarchy.html', {
            'subjects': [],
            'all_student_weaknesses': []
        })
    # Get subjects with optimized queries
    subjects = Topic.objects.all().prefetch_related(
        'general_topics',
        'general_topics__subtopics',
        'general_topics__subtopics__questions'
    )
    # Get all students' weaknesses with optimized queries
    all_student_weaknesses = []
    # Get all students (excluding admins) with their quiz levels
    students = User.objects.filter(
        is_staff=False, 
        is_superuser=False
    ).select_related('profile')
    # Get all subtopics at once to avoid multiple queries
    subtopic_ids = set()
    quiz_levels = SubjectQuizLevel.objects.select_related(
        'subject'
    ).filter(student__in=students)
    # Collect all subtopic IDs first
    for ql in quiz_levels:
        subtopic_ids.update(ql.weak_areas.get('subtopics', []))
    # Fetch all subtopics in one query
    subtopics_dict = {
        st.id: st for st in Subtopic.objects.select_related(
            'general_topic'
        ).filter(id__in=subtopic_ids)
    }
    # Process quiz levels with cached subtopics
    for ql in quiz_levels:
        student = ql.student
        for subtopic_id in ql.weak_areas.get('subtopics', []):
            subtopic = subtopics_dict.get(subtopic_id)
            if subtopic:
                all_student_weaknesses.append({
                    'student': student,
                    'subtopic': subtopic,
                    'subtopic_id': subtopic.id,  # Ensure subtopic_id is present
                    'level': ql.level,
                    'subject': ql.subject
                })
    context = {
        'subjects': subjects,
        'all_student_weaknesses': all_student_weaknesses
    }
    return render(request, 'subject_hierarchy.html', context)

@login_required
def get_subject_hierarchy(request):
    """
    Get the complete subject hierarchy including subjects, general topics, and subtopics
    """
    subjects = Topic.objects.all().prefetch_related(
        'general_topics',
        'general_topics__subtopics',
        'general_topics__subtopics__questions'
    )
    
    hierarchy = []
    for subject in subjects:
        subject_data = {
            'id': subject.id,
            'name': subject.name,
            'description': subject.description,
            'general_topics': []
        }
        
        for general_topic in subject.general_topics.all():
            topic_data = {
                'id': general_topic.id,
                'name': general_topic.name,
                'description': general_topic.description,
                'subtopics': []
            }
            
            for subtopic in general_topic.subtopics.all():
                subtopic_data = {
                    'id': subtopic.id,
                    'name': subtopic.name,
                    'description': subtopic.description,
                    'question_count': subtopic.questions.count()
                }
                topic_data['subtopics'].append(subtopic_data)
            
            subject_data['general_topics'].append(topic_data)
        
        hierarchy.append(subject_data)
    
    return JsonResponse({'hierarchy': hierarchy})

@login_required
def display_hierarchy(request):
    """
    Display the subject hierarchy in a clean, organized view
    """
    subjects = Topic.objects.all().prefetch_related(
        'general_topics',
        'general_topics__subtopics',
        'general_topics__subtopics__questions'
    )
    return render(request, 'core/display_hierarchy.html', {'subjects': subjects})

@login_required
def download_hierarchy_docx(request):
    """
    Generate and download the subject hierarchy as a docx file
    """
    # Get the subject hierarchy data
    subjects = Topic.objects.all().prefetch_related(
        'general_topics',
        'general_topics__subtopics',
        'general_topics__subtopics__questions'
    )
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Subject Hierarchy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    from datetime import datetime
    timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp.style = 'Subtitle'
    
    # Add content for each subject
    for subject in subjects:
        # Add subject heading
        subject_heading = doc.add_heading(subject.name, level=1)
        if subject.description:
            doc.add_paragraph(subject.description, style='Subtitle')
        
        # Add general topics
        for general_topic in subject.general_topics.all():
            # Add general topic heading
            topic_heading = doc.add_heading(general_topic.name, level=2)
            if general_topic.description:
                doc.add_paragraph(general_topic.description, style='Body Text')
            
            # Add subtopics
            if general_topic.subtopics.exists():
                doc.add_heading('Subtopics:', level=3)
                for subtopic in general_topic.subtopics.all():
                    subtopic_para = doc.add_paragraph(style='List Bullet')
                    subtopic_para.add_run(f'{subtopic.name}').bold = True
                    if subtopic.description:
                        subtopic_para.add_run(f' - {subtopic.description}')
                    subtopic_para.add_run(f' ({subtopic.questions.count()} questions)').italic = True
            else:
                doc.add_paragraph('No subtopics available', style='Body Text')
            
            # Add a line break between general topics
            doc.add_paragraph()
    
    # Save the document to a BytesIO object
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    # Create the HTTP response
    response = HttpResponse(
        docx_file.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=subject_hierarchy.docx'
    
    return response

@user_passes_test(is_admin)
@require_http_methods(["GET"])
def get_survey_data(request):
    """API endpoint to get survey data for dynamic updates."""
    from collections import Counter
    try:
        # Get all survey responses
        survey_responses = SurveyResponse.objects.all().order_by('-created_at')

        # Define questions and their labels for each section
        questions = {
            'experience': [
                ('experience_rating', "How would you rate your overall experience with the system?"),
                ('content_quality_rating', "How satisfied are you with the learning materials?"),
                ('quiz_quality_rating', "How helpful are the practice questions?"),
                ('system_usability_rating', "How effective is the feedback system?")
            ],
            'difficulty': [
                ('difficulty_rating', "How challenging are the practice questions?"),
                ('progress_tracking_rating', "How difficult is it to navigate the system?"),
                ('recommendation_quality_rating', "How complex are the learning materials?"),
                ('study_materials_rating', "How challenging is the assessment process?")
            ],
            'control': [
                ('peu_control', "How much control do you have over your learning pace?"),
                ('knowledge_improvement', "How easy is it to track your progress?"),
                ('confidence_improvement', "How well can you customize your learning experience?"),
                ('study_habits_improvement', "How much control do you have over the content you study?")
            ]
        }

        question_tallies = {}
        for section, qs in questions.items():
            question_tallies[section] = []
            for field, label in qs:
                responses = survey_responses.values_list(field, flat=True)
                tally = Counter(responses)
                # Ensure keys are integers 1-5
                tally_dict = {i: tally.get(i, 0) for i in range(1, 6)}
                total = sum(tally_dict.values())
                weighted_mean = round(sum(i * tally_dict[i] for i in range(1, 6)) / total, 2) if total else 0
                question_tallies[section].append({
                    'label': label,
                    'tally': tally_dict,
                    'weighted_mean': weighted_mean,
                    'total': total
                })

        # Ensure all sections are present
        for section in ['experience', 'difficulty', 'control']:
            if section not in question_tallies:
                question_tallies[section] = []

        # Process survey responses for detailed table
        responses_data = []
        for response in survey_responses:
            response_data = {
                'student': response.student.username,
                'experience_rating': response.experience_rating,
                'difficulty_rating': response.difficulty_rating,
                'peu_control': response.peu_control,
                'feedback': response.suggestions,
                'created_at': response.created_at.strftime('%Y-%m-%d %H:%M:%S')
            }
            responses_data.append(response_data)

        # Calculate averages
        total_responses = survey_responses.count()
        avg_experience = survey_responses.aggregate(avg=Avg('experience_rating'))['avg'] or 0
        avg_difficulty = survey_responses.aggregate(avg=Avg('difficulty_rating'))['avg'] or 0
        avg_control = survey_responses.aggregate(avg=Avg('peu_control'))['avg'] or 0

        return JsonResponse({
            'survey_responses': responses_data,
            'question_tallies': question_tallies,
            'averages': {
                'experience': round(avg_experience, 1),
                'difficulty': round(avg_difficulty, 1),
                'control': round(avg_control, 1)
            },
            'total_responses': total_responses
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@require_GET
def student_weak_topics_ajax(request, student_id):
    try:
        student = User.objects.get(id=student_id)
    except User.DoesNotExist:
        return JsonResponse({'error': 'Student not found'}, status=404)
    quiz_levels = SubjectQuizLevel.objects.filter(student=student)
    weak_topics = []
    for ql in quiz_levels:
        # Get all weak subtopics for this subject and level
        weak_subtopic_ids = ql.weak_areas.get('subtopics', [])
        if not weak_subtopic_ids:
            continue
        # Get all subtopics for this subject and level
        all_subtopics = list(Subtopic.objects.filter(general_topic__subject=ql.subject))
        num_subtopics = len(all_subtopics) if all_subtopics else 1
        # Get required questions from settings
        try:
            settings = QuestionGenerationSetting.objects.get(level=ql.level)
            total_required = settings.questions_per_topic
        except QuestionGenerationSetting.DoesNotExist:
            total_required = 10
        base_required = total_required // num_subtopics
        remainder = total_required % num_subtopics
        # Assign required per subtopic
        required_per_subtopic = {}
        for idx, subtopic in enumerate(all_subtopics):
            required_per_subtopic[subtopic.id] = base_required + (1 if idx < remainder else 0)
        # For each weak subtopic, get count and lacking
        for subtopic_id in weak_subtopic_ids:
            try:
                subtopic = Subtopic.objects.get(id=subtopic_id)
                count = Question.objects.filter(subtopic=subtopic, level=ql.level).count()
                required = required_per_subtopic.get(subtopic.id, base_required)
                lacking = max(0, required - count)
                weak_topics.append({
                    'subject': ql.subject.name,
                    'subtopic': subtopic.name,
                    'level': ql.level,
                    'description': subtopic.description or '',
                    'subtopic_id': subtopic.id,
                    'required': required,
                    'count': count,
                    'lacking': lacking
                })
            except Subtopic.DoesNotExist:
                continue
    return JsonResponse({'weak_topics': weak_topics})

@login_required
def question_count(request):
    """API endpoint to get the count of available questions for each subtopic in a subject at a specific level."""
    subject = request.GET.get('subject')
    level = request.GET.get('level')
    student_id = request.GET.get('student_id')
    
    if not subject or not level or not student_id:
        return JsonResponse({'error': 'Subject, level, and student_id are required'}, status=400)
    
    try:
        # Get the level object
        level_obj = Level.objects.get(id=level)
        
        # Get the question generation settings for this level
        try:
            settings = QuestionGenerationSetting.objects.get(level=level_obj)
            total_required_questions = settings.questions_per_topic
            difficulty_distribution = {
                'easy': settings.easy_percentage,
                'medium': settings.medium_percentage,
                'hard': settings.hard_percentage
            }
        except QuestionGenerationSetting.DoesNotExist:
            total_required_questions = 10  # Default value if no settings exist
            difficulty_distribution = {
                'easy': 30,
                'medium': 50,
                'hard': 20
            }
        
        # Get student's weak areas for this subject and level
        try:
            quiz_level = SubjectQuizLevel.objects.get(
                student_id=student_id,
                subject__name=subject,
                level=level
            )
            weak_subtopic_ids = quiz_level.weak_areas.get('subtopics', [])
        except SubjectQuizLevel.DoesNotExist:
            return JsonResponse({'error': 'Student quiz level not found'}, status=404)
        
        # Get all questions for the subject at the specified level
        questions = Question.objects.filter(
            subtopic__general_topic__subject__name=subject,
            level=level
        ).select_related('subtopic')
        
        # Group questions by subtopic and count them
        subtopic_counts = {}
        for question in questions:
            subtopic = question.subtopic.name
            if subtopic not in subtopic_counts:
                subtopic_counts[subtopic] = 0
            subtopic_counts[subtopic] += 1
        
        # Calculate required questions per subtopic based on settings
        num_weak_subtopics = len(weak_subtopic_ids)
        if num_weak_subtopics == 0:
            return JsonResponse({'error': 'No weak subtopics found'}, status=404)
            
        # Calculate base questions per subtopic and remainder
        base_questions = total_required_questions // num_weak_subtopics
        remainder = total_required_questions % num_weak_subtopics
        
        # Only include weak subtopics in the response
        topics = []
        for idx, subtopic_id in enumerate(weak_subtopic_ids):
            try:
                subtopic = Subtopic.objects.get(id=subtopic_id)
                count = subtopic_counts.get(subtopic.name, 0)
                # Distribute remainder questions to first few subtopics
                required = base_questions + (1 if idx < remainder else 0)
                topics.append({
                    'subtopic': subtopic.name,
                    'subtopic_id': subtopic.id,  # Add the subtopic ID
                    'count': count,
                    'required': required,
                    'lacking': max(0, required - count)
                })
            except Subtopic.DoesNotExist:
                continue
        
        return JsonResponse({
            'topics': topics,
            'required_questions': total_required_questions,
            'difficulty_distribution': difficulty_distribution
        })
        
    except Level.DoesNotExist:
        return JsonResponse({'error': 'Level not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_http_methods(['POST'])
@csrf_exempt
def generate_questions(request):
    """API endpoint to generate questions for a specific subtopic using Ollama."""
    try:
        print("Starting question generation process...")
        data = json.loads(request.body)
        subtopic_id = data.get('subtopic_id')
        level = data.get('level')
        student_id = data.get('student_id')
        
        print(f"Received request for subtopic_id: {subtopic_id}, level: {level}, student_id: {student_id}")
        
        if not all([subtopic_id, level, student_id]):
            print("Missing required parameters")
            return JsonResponse({'error': 'Missing required parameters'}, status=400)
        
        # Get the subtopic
        subtopic = get_object_or_404(Subtopic, id=subtopic_id)
        print(f"Found subtopic: {subtopic.name}")
        
        # Get the student
        student = get_object_or_404(User, id=student_id)
        print(f"Found student: {student.username}")
        
        # Get the level settings
        level_obj = Level.objects.get(id=level)
        settings = QuestionGenerationSetting.objects.get(level=level_obj)
        print(f"Found level settings for level {level}")
        
        # Calculate total required questions for this level
        total_required_questions = settings.questions_per_topic
        print(f"Total required questions for level: {total_required_questions}")
        
        # Get all weak subtopics for this student in this subject and level
        try:
            quiz_level = SubjectQuizLevel.objects.get(
                student=student,
                subject=subtopic.general_topic.subject,
                level=level
            )
            weak_subtopic_ids = quiz_level.weak_areas.get('subtopics', [])
            print(f"Found {len(weak_subtopic_ids)} weak subtopics for student")
        except SubjectQuizLevel.DoesNotExist:
            print("Student quiz level not found")
            return JsonResponse({'error': 'Student quiz level not found'}, status=404)
        
        # Count existing questions for each weak subtopic
        subtopic_counts = {}
        for weak_id in weak_subtopic_ids:
            count = Question.objects.filter(
                subtopic_id=weak_id,
                level=level
            ).count()
            subtopic_counts[weak_id] = count
            print(f"Subtopic {weak_id} has {count} existing questions")
        
        # Calculate base questions per subtopic and remainder
        num_weak_subtopics = len(weak_subtopic_ids)
        if num_weak_subtopics == 0:
            print("No weak subtopics found")
            return JsonResponse({'error': 'No weak subtopics found'}, status=404)
            
        base_questions = total_required_questions // num_weak_subtopics
        remainder = total_required_questions % num_weak_subtopics
        print(f"Base questions per subtopic: {base_questions}, Remainder: {remainder}")
        
        # Calculate questions lacking for each subtopic
        questions_lacking = {}
        for idx, weak_id in enumerate(weak_subtopic_ids):
            required = base_questions + (1 if idx < remainder else 0)
            existing = subtopic_counts.get(weak_id, 0)
            questions_lacking[weak_id] = max(0, required - existing)
            print(f"Subtopic {weak_id} needs {questions_lacking[weak_id]} more questions")
        
        # Get questions lacking for the current subtopic
        questions_to_generate = questions_lacking.get(subtopic_id, 0)
        print(f"Need to generate {questions_to_generate} questions for subtopic {subtopic_id}")
        
        if questions_to_generate <= 0:
            print(f"No questions needed for {subtopic.name}")
            return JsonResponse({
                'status': 'success',
                'generated_count': 0,
                'message': f'No questions needed for {subtopic.name}'
            })

        # Calculate question distribution based on settings
        easy_count = int(questions_to_generate * (settings.easy_percentage / 100))
        medium_count = int(questions_to_generate * (settings.medium_percentage / 100))
        hard_count = questions_to_generate - easy_count - medium_count
        
        print(f"Question distribution - Easy: {easy_count}, Medium: {medium_count}, Hard: {hard_count}")
        
        # Get question types from settings
        question_types = settings.question_types
        if not question_types:
            print("No question types specified in settings")
            return JsonResponse({'error': 'No question types specified in settings'}, status=400)
            
        # Calculate questions per type
        type_distribution = {}
        questions_per_type = questions_to_generate // len(question_types)
        remaining_questions = questions_to_generate % len(question_types)
        
        for q_type in question_types:
            type_distribution[q_type] = questions_per_type
        if remaining_questions > 0:
            type_distribution[question_types[0]] += remaining_questions
            
        print(f"Question type distribution: {type_distribution}")
        
        # Prepare the prompt for Ollama
        prompt = f"""
Generate {questions_to_generate} questions for the subtopic '{subtopic.name}' under the topic '{subtopic.general_topic.name}'.
The questions should be appropriate for level {level}.

Question Distribution:
-- {easy_count} easy questions
-- {medium_count} medium questions
-- {hard_count} hard questions

Question Types Needed:
{chr(10).join([f"- {q_type}: {count} questions" for q_type, count in type_distribution.items()])}

Format each question exactly as JSON with these fields:
-- question_type: must be one of {question_types}
-- question_text: The question
-- difficulty: 'easy', 'medium', or 'hard'
-- choices: List of 4 answer choices (for multiple_choice) or ["True", "False"] (for true_false)
-- correct_answer: The correct answer
-- points: Difficulty points (1-5)

Return only a JSON array of questions, no extra text.
"""
        print("Starting Ollama question generation...")
        print(f"Prompt: {prompt}")
        
        # Check if Ollama is running
        try:
            health_check = requests.get('http://localhost:11434/api/tags', timeout=5)
            health_check.raise_for_status()
            print("Ollama server is running")
        except requests.exceptions.RequestException as e:
            print(f"Ollama server is not running or not accessible: {str(e)}")
            return JsonResponse({
                'error': 'Ollama server is not running. Please start Ollama and try again.',
                'details': str(e)
            }, status=503)
        
        # Make API call to Ollama
        try:
            print("Making request to Ollama API...")
            response = requests.post(
                'http://localhost:11434/api/generate',
                json={
                    'model': 'mistral',
                    'prompt': prompt,
                    'stream': False,
                    'options': {
                        'temperature': 0.7,
                        'max_tokens': 1000,
                        'num_ctx': 4096,
                        'num_thread': 4,
                        'num_gpu': 1
                    }
                },
                timeout=600  # 10 minutes timeout
            )
            print(f"Ollama API response status: {response.status_code}")
            response.raise_for_status()
            generated_text = response.json().get('response', '')
            print("Received response from Ollama")
            print(f"Generated text length: {len(generated_text)}")
            
            try:
                questions = json.loads(generated_text)
                print(f"Successfully parsed {len(questions)} questions from response")
            except json.JSONDecodeError as e:
                print(f"Failed to parse generated questions as JSON: {str(e)}")
                print(f"Generated text: {generated_text[:500]}...")  # Print first 500 chars for debugging
                return JsonResponse({'error': 'Failed to parse generated questions'}, status=500)
            
            # Create Question objects
            generated_questions = []
            current_school_year = SchoolYear.objects.order_by('-start_date').first()
            
            for q in questions:
                try:
                    q_text = q.get('question_text', '').strip()
                    q_type = q.get('question_type', '').strip().lower()
                    q_points = int(q.get('points', 1))
                    q_correct = q.get('correct_answer', '').strip()
                    q_choices = q.get('choices', [])
                    q_difficulty = q.get('difficulty', 'medium').lower()
                    
                    if q_type not in question_types:
                        print(f"Skipping question with invalid type: {q_type}")
                        continue
                        
                    difficulty_points = {'easy': 1, 'medium': 3, 'hard': 5}
                    q_points = difficulty_points.get(q_difficulty, 1)
                    
                    if q_type == 'multiple_choice':
                        question = Question.objects.create(
                            question_text=q_text,
                            subtopic=subtopic,
                            points=q_points,
                            correct_answer=q_correct,
                            question_type='multiple_choice',
                            level=level,
                            school_year=current_school_year
                        )
                        for choice_text in q_choices:
                            if isinstance(choice_text, list):
                                choice_text = ' '.join(str(x) for x in choice_text)
                            Choice.objects.create(
                                question=question,
                                choice_text=str(choice_text),
                                is_correct=(str(choice_text) == q_correct)
                            )
                        generated_questions.append(question)
                        print(f"Created multiple choice question: {q_text[:50]}...")
                        
                    elif q_type == 'true_false':
                        question = Question.objects.create(
                            question_text=q_text,
                            subtopic=subtopic,
                            points=q_points,
                            correct_answer=q_correct,
                            question_type='true_false',
                            level=level,
                            school_year=current_school_year
                        )
                        for choice_text in ['True', 'False']:
                            Choice.objects.create(
                                question=question,
                                choice_text=choice_text,
                                is_correct=(choice_text == q_correct)
                            )
                        generated_questions.append(question)
                        print(f"Created true/false question: {q_text[:50]}...")
                        
                    elif q_type in ['short_answer', 'essay']:
                        question = Question.objects.create(
                            question_text=q_text,
                            subtopic=subtopic,
                            points=q_points,
                            correct_answer=q_correct,
                            question_type=q_type,
                            level=level,
                            school_year=current_school_year
                        )
                        generated_questions.append(question)
                        print(f"Created {q_type} question: {q_text[:50]}...")
                        
                except Exception as e:
                    print(f"Error creating question: {str(e)}")
                    continue
            
            print(f"Successfully generated {len(generated_questions)} questions")
            return JsonResponse({
                'status': 'success',
                'generated_count': len(generated_questions),
                'message': f'Successfully generated {len(generated_questions)} questions for {subtopic.name}',
                'questions_lacking': questions_to_generate
            })
            
        except requests.exceptions.RequestException as e:
            print(f"Error making API request to Ollama: {str(e)}")
            return JsonResponse({
                'error': 'Failed to generate questions with Ollama',
                'details': str(e)
            }, status=500)
        
    except json.JSONDecodeError:
        print("Invalid JSON data received")
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        print(f"Error in generate_questions: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_http_methods(['POST'])
@csrf_exempt
def add_manual_question(request):
    """
    API endpoint to add a question manually.
    """
    try:
        data = json.loads(request.body)
        subtopic_id = data.get('subtopic_id')
        level = data.get('level')
        student_id = data.get('student_id')
        question_text = data.get('question_text')
        question_type = data.get('question_type')
        difficulty = data.get('difficulty')
        options = data.get('options', [])
        correct_answer = data.get('correct_answer')
        explanation = data.get('explanation', '')

        # Validate required fields
        if not all([subtopic_id, level, student_id, question_text, question_type, difficulty, correct_answer]):
            return JsonResponse({
                'error': 'Missing required fields'
            }, status=400)

        # Get the subtopic
        try:
            subtopic = Subtopic.objects.get(id=subtopic_id)
        except Subtopic.DoesNotExist:
            return JsonResponse({
                'error': 'Subtopic not found'
            }, status=404)

        # Create the question
        question = Question.objects.create(
            subtopic=subtopic,
            level=level,
            question_text=question_text,
            question_type=question_type,
            difficulty=difficulty,
            correct_answer=correct_answer,
            explanation=explanation
        )

        # Add options for multiple choice questions
        if question_type == 'multiple_choice' and options:
            for option in options:
                QuestionOption.objects.create(
                    question=question,
                    option_text=option,
                    is_correct=(option == correct_answer)
                )

        return JsonResponse({
            'message': 'Question added successfully',
            'question_id': question.id
        })

    except json.JSONDecodeError:
        return JsonResponse({
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'error': str(e)
        }, status=500)

@login_required
@api_view(['GET'])
def get_subtopic_data(request, subtopic_id):
    """
    Get data for a specific subtopic including its subject, general topic, and level information.
    """
    try:
        print(f"Fetching data for subtopic_id: {subtopic_id}")  # Debug log
        
        if not subtopic_id or subtopic_id == 'undefined':
            return JsonResponse({'error': 'Invalid subtopic ID'}, status=400)
            
        subtopic = get_object_or_404(Subtopic, id=subtopic_id)
        subject = subtopic.general_topic.subject
        
        # Get school year if available
        school_year_id = subject.school_year.id if hasattr(subject, 'school_year') and subject.school_year else None
        
        # Get the most recent question for this subtopic to infer the level
        question = subtopic.questions.order_by('-level').first()
        level = question.level if question and hasattr(question, 'level') else 1
        
        data = {
            'school_year': school_year_id,
            'subject': subject.id,
            'general_topic': subtopic.general_topic.id,
            'level': level
        }
        
        print(f"Returning data: {data}")  # Debug log
        return JsonResponse(data)
        
    except Subtopic.DoesNotExist:
        print(f"Subtopic not found with id: {subtopic_id}")  # Debug log
        return JsonResponse({'error': 'Subtopic not found'}, status=404)
    except Exception as e:
        print(f"Error in get_subtopic_data: {str(e)}")  # Debug log
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def get_school_years(request):
    from .models import SchoolYear
    return JsonResponse(list(SchoolYear.objects.values('id', 'name')), safe=False)

@user_passes_test(is_admin)
def change_admin_credentials(request):
    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        username = request.POST.get('username')
        email = request.POST.get('email')
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')
        
        # Update first name and last name
        if first_name:
            request.user.first_name = first_name
        if last_name:
            request.user.last_name = last_name
        
        # Validate username
        if username and username != request.user.username:
            if User.objects.filter(username=username).exclude(id=request.user.id).exists():
                messages.error(request, 'Username already exists.')
                return redirect('change_admin_credentials')
            request.user.username = username
        
        # Validate email
        if email and email != request.user.email:
            if User.objects.filter(email=email).exclude(id=request.user.id).exists():
                messages.error(request, 'Email already exists.')
                return redirect('change_admin_credentials')
            request.user.email = email
        
        # Validate password
        if new_password:
            if new_password != confirm_password:
                messages.error(request, 'Passwords do not match.')
                return redirect('change_admin_credentials')
            request.user.set_password(new_password)
        
        try:
            request.user.save()
            messages.success(request, 'Your credentials have been updated successfully!')
            return redirect('admin_dashboard')
        except Exception as e:
            messages.error(request, f'Error updating credentials: {str(e)}')
    
    if request.user.is_authenticated:
        request.session['last_template_access'] = str(timezone.now())
    return render(request, 'admin/change_admin_credentials.html', {
        'user': request.user
    })

@staff_member_required
def performance_metrics_view(request):
    metrics = PerformanceMetric.objects.all().order_by('-created_at')
    return render(request, 'core/performance_metrics.html', {'metrics': metrics})
