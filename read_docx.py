from docx import Document
import os

def read_docx(file_path):
    try:
        # Load the document
        doc = Document(file_path)
        
        # Print the full text
        print("\n=== Document Content ===\n")
        for paragraph in doc.paragraphs:
            print(paragraph.text)
            
        # Print tables if any
        if doc.tables:
            print("\n=== Tables ===\n")
            for i, table in enumerate(doc.tables, 1):
                print(f"\nTable {i}:")
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    print(" | ".join(row_text))
                    
    except Exception as e:
        print(f"Error reading the document: {str(e)}")

if __name__ == "__main__":
    # Path to the DOCX file
    docx_path = os.path.join("System_paper", "galicia-buang.docx")
    
    # Check if file exists
    if os.path.exists(docx_path):
        read_docx(docx_path)
    else:
        print(f"File not found: {docx_path}") 